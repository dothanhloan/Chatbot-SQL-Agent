import os
import uuid
import requests
from typing import Union, List, Dict, Any
from dotenv import load_dotenv

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document

# LangChain - OpenAI
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Static DOCX RAG (simple, like old CLI bot)
from static_data_rag import initialize_rag, get_rag_instance, build_static_context

# ==========================================================
# 1. SETUP & CẤU HÌNH
# ==========================================================
load_dotenv()

# BẮT BUỘC phải có OpenAI API Key
if not os.environ.get("OPENAI_API_KEY"):
    raise RuntimeError("❌ Chưa cấu hình OPENAI_API_KEY")

HRM_API_URL = "https://hrm.icss.com.vn/ICSS/api/execute-sql"

app = FastAPI(title="ICS HRM SQL Chatbot API", version="3.1 - OpenAI + Static DOCX RAG (simple)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==========================================================
# STARTUP: initialize simple static RAG (DOCX)
# ==========================================================
@app.on_event("startup")
async def startup_event() -> None:
    try:
        ok = initialize_rag()
        if ok and get_rag_instance().is_initialized():
            print("[StaticRAG] Initialized successfully")
        else:
            print("[StaticRAG] Initialization skipped or failed (no DOCX/data)")
    except Exception as e:  # noqa: BLE001
        print(f"[StaticRAG] Initialization error: {e}")

# Tạo thư mục lưu file tạm
EXPORT_DIR = "./static/reports"
if not os.path.exists(EXPORT_DIR):
    os.makedirs(EXPORT_DIR)

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_word_report(data, title="BÁO CÁO HRM", filename_prefix="report", question="", summary=""):
    """Sinh file .docx từ dữ liệu SQL - Định dạng báo cáo khoa học"""
    if not data: return None
    
    # Đảm bảo data là list
    if isinstance(data, dict):
        data = [data]
    
    # 1. Khởi tạo file Word
    doc = Document()
    
    # === PHẦN TIÊU ĐỀ ===
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm đường kẻ và thông tin thời gian
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Khoảng trống
    
    # === PHẦN CÂU HỎI ===
    if question:
        doc.add_heading("1. Yêu cầu truy vấn", level=1)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f'"{question}"')
        q_run.font.italic = True
        q_run.font.size = Pt(11)
        doc.add_paragraph()
    
    # === PHẦN TÓM TẮT KẾT QUẢ ===
    if summary:
        doc.add_heading("2. Tóm tắt kết quả", level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.space_after = Pt(12)
        doc.add_paragraph()
    
    # === PHẦN BẢNG DỮ LIỆU CHI TIẾT ===
    section_num = 3 if question and summary else (2 if question or summary else 1)
    doc.add_heading(f"{section_num}. Dữ liệu chi tiết ({len(data)} bản ghi)", level=1)
    
    # Lấy headers từ keys của dòng đầu tiên
    headers = list(data[0].keys())
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ghi header với định dạng đẹp
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h).upper().replace('_', ' ')
        # Định dạng header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        
    # Ghi dữ liệu
    for item in data:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            cell_value = item.get(h, '')
            row_cells[i].text = str(cell_value) if cell_value is not None else ''
            # Định dạng cell
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # === PHẦN FOOTER ===
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("─" * 50)
    footer_run.font.color.rgb = RGBColor(200, 200, 200)
    
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = footer_info.add_run("Báo cáo được tạo tự động bởi ICS HRM Chatbot")
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
            
    # 3. Lưu file
    filename = f"{filename_prefix}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    
    return filepath

def create_pdf_report(data, title="BAO CAO HRM", filename_prefix="report"):
    """Sinh file .pdf từ dữ liệu SQL"""
    if not data: return None

    pdf = FPDF()
    pdf.add_page()
    
    # Lưu ý: FPDF mặc định không hỗ trợ tiếng Việt Unicode tốt trừ khi add font ngoài.
    # Ở đây ta dùng font chuẩn Arial (sẽ mất dấu tiếng Việt nếu không config thêm font)
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt=title, ln=1, align='C')
    
    # Ghi dữ liệu dòng
    for item in data:
        row_str = " | ".join([f"{str(v)}" for k,v in item.items()])
        # Encode để tránh lỗi ký tự lạ
        safe_str = row_str.encode('latin-1', 'replace').decode('latin-1') 
        pdf.cell(0, 10, txt=safe_str, ln=1)
        
    filename = f"{filename_prefix}_{uuid.uuid4().hex[:6]}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)
    pdf.output(filepath)
    
    return filepath

# ==========================================================
# 2. SCHEMA REQUEST / RESPONSE
# ==========================================================
class ChatRequest(BaseModel):
    question: str


class ChatResponse(BaseModel):
    sql: Union[str, None]
    data: Union[List, Dict, Any, None]
    answer: str
    download_url: Union[str, None] = None


# ==========================================================
# 3. KHỞI TẠO LLM (OPENAI)
# ==========================================================
# LLM cho sinh SQL (cần token ít hơn)
llm_sql = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    max_tokens=1000   # Đủ cho SQL query
)

# LLM cho trả lời (cần token nhiều hơn để liệt kê danh sách dài)
llm_answer = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    max_tokens=4000   # Tăng để hỗ trợ danh sách dài (26+ dòng)
)

# LLM router để phân loại câu hỏi HRM_SQL vs STATIC_DOC
llm_router = ChatOpenAI(
  model="gpt-4o-mini",
  temperature=0,
  max_tokens=10,
)

# Giữ lại llm tham chiếu để tương thích với code cũ
llm = llm_sql
# ==========================================================
# 2. SCHEMA & LUẬT NGHIỆP VỤ (Nguồn: HRM_SCHEMA.docx)
# ==========================================================
HRM_SCHEMA_RAW = """
-- CHẤM CÔNG [Source: 7] --
BẢNG cham_cong: id, nhan_vien_id, ngay (date), check_in (time), check_out (time).

-- NHÂN SỰ [Source: 12] --
BẢNG nhanvien: id, ho_ten, email, so_dien_thoai, phong_ban_id, chuc_vu, vai_tro, luong_co_ban, trang_thai_lam_viec, ngay_vao_lam, gioi_tinh.
BẢNG phong_ban: id, ten_phong, truong_phong_id [Source: 13].

-- LƯƠNG & KPI [Source: 10, 11] --
BẢNG luong: id, nhan_vien_id, thang, nam, luong_co_ban, phu_cap, khoan_tru.
BẢNG luu_kpi: id, nhan_vien_id, thang, nam, diem_kpi, xep_loai.
BẢNG ngay_phep_nam: id, nhan_vien_id, nam, tong_ngay_phep, ngay_phep_con_lai.

-- DỰ ÁN & CÔNG VIỆC [Source: 7, 8, 9] --
BẢNG du_an: id, ten_du_an, lead_id (PM), phong_ban (varchar), trang_thai_duan, ngay_bat_dau, ngay_ket_thuc.
BẢNG cong_viec: id, ten_cong_viec, nguoi_giao_id, han_hoan_thanh, trang_thai, muc_do_uu_tien, du_an_id.
BẢNG cong_viec_nguoi_nhan: id, cong_viec_id, nhan_vien_id.
BẢNG cong_viec_tien_do: id, cong_viec_id, phan_tram.
BẢNG cong_viec_quy_trinh: id, cong_viec_id, ten_buoc, mo_ta, trang_thai, ngay_bat_dau, ngay_ket_thuc, ngay_tao.

-- HỆ THỐNG --
BẢNG cau_hinh_he_thong: id, ten_cau_hinh, gia_tri, mo_ta, ngay_tao.

-- TÀI LIỆU & HỆ THỐNG [Source: 14] --
BẢNG tai_lieu: id, ten_tai_lieu, mo_ta, file_path, file_name, file_type, file_size, loai_tai_lieu, trang_thai, doi_tuong_xem, luot_xem, luot_tai, ngay_tao, ngay_cap_nhat, nguoi_tao_id.
BẢNG thong_bao: id, tieu_de, noi_dung, nguoi_nhan_id.
BẢNG lich_trinh: id, tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc, ngay_tao.

-- QUYỀN HẠN & PHÂN QUYỀN [Source: 15] --
BẢNG quyen: id, nhom_quyen, ma_quyen, ten_quyen.

"""

# Kết hợp Schema thô với Luật nghiệp vụ (Enhanced Schema)
HRM_SCHEMA_ENHANCED = f"""
DANH SÁCH BẢNG VÀ LUẬT NGHIỆP VỤ BẮT BUỘC (DATA TRUTH):

1. **QUY TẮC ĐI MUỘN (08:06 RULE) - BẮT BUỘC:**
   - Định nghĩa: Nhân viên CÓ đi làm (check_in NOT NULL) nhưng giờ vào **từ 08:06:00 trở đi**.
   - SQL Logic: `check_in >= '08:06:00'`.
   - LƯU Ý: Tuyệt đối CẤM dùng `> 08:05`.
   - Phân biệt: Nếu không có dữ liệu chấm công -> Là Vắng mặt (Absent), dùng `NOT IN`.

2. **BẢNG `phong_ban` & `du_an`:**
   - Tìm tên phòng ban: BẮT BUỘC dùng `LIKE` (VD: `LIKE '%Marketing%'`). **CẤM** dùng `=`.
   - Dự án của phòng: Cột `phong_ban` trong bảng `du_an` là text (varchar). Tìm dự án theo phòng phải query trên bảng `du_an` (dùng LIKE).
   - **PHÂN BIỆT: DỰ ÁN CHUNG vs DỰ ÁN ĐANG CHẠY (RẤT QUAN TRỌNG):**
     + **Câu hỏi: "Dự án của phòng Marketing?"** (chung chung)
       → Lọc bỏ Tạm Ngưng: `WHERE phong_ban LIKE '%Marketing%' AND trang_thai_duan != 'Tạm Ngưng'`
     + **Câu hỏi: "Dự án đang thực hiện của phòng Marketing?"** hoặc **"Trạng thái đang chạy theo phòng ban?"** (rõ ý)
       → Chỉ lấy đang thực hiện: `WHERE phong_ban LIKE '%Marketing%' AND trang_thai_duan = 'Đang thực hiện'`
     + **Câu hỏi: "Liệt kê theo phòng ban với trạng thái"**
       → Cần `GROUP BY phong_ban, trang_thai_duan` hoặc chỉ `GROUP BY phong_ban` nếu hỏi số lượng dự án đang chạy
   - **PHÂN BIỆT LIỆT KÊ (List) vs ĐẾM (Count) - RẤT QUAN TRỌNG:**
     + **Câu hỏi: "Liệt kê dự án đang thực hiện theo phòng ban"** (liệt kê chi tiết)
       → Dùng `SELECT phong_ban, ten_du_an ... ORDER BY phong_ban`
       → **KHÔNG dùng COUNT hay GROUP BY** (sẽ mất chi tiết)
       → SQL: `SELECT phong_ban, ten_du_an FROM du_an WHERE trang_thai_duan = 'Đang thực hiện' ORDER BY phong_ban`
     + **Câu hỏi: "Bao nhiêu dự án đang thực hiện ở từng phòng ban?"** (đếm số lượng)
       → Dùng `SELECT phong_ban, COUNT(id) ... GROUP BY phong_ban`
       → SQL: `SELECT phong_ban, COUNT(id) as so_du_an FROM du_an WHERE trang_thai_duan = 'Đang thực hiện' GROUP BY phong_ban`
     + **TỪKHÓA PHÂN BIỆT:**
       → "Liệt kê", "Danh sách", "Tìm", "Cho tôi xem" → Dùng SELECT chi tiết, không COUNT
       → "Bao nhiêu", "Tổng", "Đếm", "Thống kê", "Số lượng" → Dùng COUNT + GROUP BY

3. **BẢNG `cong_viec` (Task) - PHÂN BIỆT CÔNG VIỆC vs DỰ ÁN (RẤT QUAN TRỌNG):**
   - **CÔNG VIỆC (cong_viec):** Có cột `trang_thai` (trạng thái công việc), `ten_cong_viec`, `han_hoan_thanh`
   - **DỰ ÁN (du_an):** Có cột `trang_thai_duan` (trạng thái dự án), `ten_du_an`, `phong_ban` (text)
   - **Câu hỏi: "Công việc nào đang thực hiện của phòng kỹ thuật?"** (hỏi CÔNG VIỆC)
     + Phải JOIN `cong_viec` → `cong_viec_nguoi_nhan` → `nhanvien` (để lấy phòng ban của người làm)
     + HOẶC: JOIN `cong_viec` → `du_an` (để lấy phòng ban từ dự án)
     + SQL: `SELECT DISTINCT cv.ten_cong_viec FROM cong_viec cv JOIN du_an d ON cv.du_an_id = d.id WHERE cv.trang_thai = 'Đang thực hiện' AND d.phong_ban LIKE '%kỹ thuật%'`
   - **Câu hỏi: "Dự án nào của phòng kỹ thuật?"** (hỏi DỰ ÁN)
     + Query trực tiếp bảng `du_an`
     + SQL: `SELECT ten_du_an FROM du_an WHERE phong_ban LIKE '%kỹ thuật%'`
   - **⚠️ CẢNH BÁO:** KHÔNG được query công việc từ bảng `du_an`, hoặc query dự án từ `cong_viec` mà không JOIN đúng cách
   - Muốn biết ai thực hiện công việc -> Phải JOIN bảng `cong_viec_nguoi_nhan`.
   - Trễ hạn: `han_hoan_thanh < CURRENT_DATE` AND `trang_thai != 'Hoàn thành'`.

4. **LUẬT TRA CỨU LƯƠNG (QUAN TRỌNG - SỬA ĐỔI):**
   - Bảng `luong` hiện tại KHÔNG có dữ liệu.
   - Khi người dùng hỏi về Lương (cơ bản, thu nhập...), **HÃY TRUY VẤN TỪ BẢNG `nhanvien`**.
   - Cột cần lấy: `nhanvien.luong_co_ban`.
   - Tuyệt đối không JOIN bảng `luong`.

5. **LUẬT DỰ ÁN & CÔNG VIỆC (QUAN TRỌNG):**
   - **Tìm Dự án theo phòng:** Cột `du_an.phong_ban` là text -> Dùng `LIKE`, CẤM JOIN bảng `phong_ban`.
   - **Tìm Quản lý (PM/Lead):** 
     + Cột `lead_id` trong `du_an` chỉ là số.
     + BẮT BUỘC JOIN bảng `nhanvien`: `ON du_an.lead_id = nhanvien.id`.
     + SELECT `nhanvien.ho_ten`.
   - **Người thực hiện task:** JOIN `cong_viec` -> `cong_viec_nguoi_nhan` -> `nhanvien`.
   - **LỌC DỰ ÁN TẠMDỪNG (CRITICAL):**
     + **Câu hỏi: "Có bao nhiêu dự án?"** (chung chung, không rõ ý)
       -> BẮT BUỘC lọc bỏ dự án tạm dừng: `WHERE trang_thai_duan != 'Tạm Ngưng'`
       -> Lý do: Người dùng thường chỉ quan tâm dự án hoạt động, không phải dự án bị ngưng
     + **Câu hỏi: "Danh sách tất cả dự án?"** (bao gồm tạm dừng)
       -> KHÔNG lọc, return tất cả dự án
     + **Câu hỏi: "Dự án nào tạm dừng?"** hoặc **"Dự án bị ngưng?"** (rõ ý)
       -> KHÔNG lọc, hoặc lọc CHỈ những cái tạm dừng
     + **Giá trị trạng thái dự án:** Giá trị chính xác là 'Tạm Ngưng' (để lọc dự án tạm dừng), các trạng thái khác: 'Đang thực hiện', 'Hoàn thành', v.v.
   - **PHÂN BIỆT NGÀY BẮT ĐẦU VS NGÀY KẾT THÚC (CRITICAL):**
     + **Câu hỏi: "Dự án nào được bắt đầu gần đây nhất?"** → Dùng `ngay_bat_dau`: `ORDER BY ngay_bat_dau DESC LIMIT 1`
     + **Câu hỏi: "Dự án nào kết thúc gần đây nhất?"** → Dùng `ngay_ket_thuc`: `ORDER BY ngay_ket_thuc DESC LIMIT 1`
     + **Câu hỏi: "Dự án nào bắt đầu lâu nhất?"** → Dùng `ngay_bat_dau`: `ORDER BY ngay_bat_dau ASC LIMIT 1`
     + **Câu hỏi: "Dự án nào kết thúc lâu nhất?"** → Dùng `ngay_ket_thuc`: `ORDER BY ngay_ket_thuc ASC LIMIT 1`
     + **⚠️ CẢNH BÁO:** CẤM nhầm lẫn giữa "bắt đầu" (ngay_bat_dau) và "kết thúc" (ngay_ket_thuc)

5.1. **LUẬT TRA CỨU PHÒNG BAN CỦA NHÂN VIÊN (RẤT QUAN TRỌNG):**
   - Bảng `nhanvien` KHÔNG có cột `phong_ban` dạng text, chỉ có `phong_ban_id`.
   - Khi người dùng hỏi: "Nhân viên X thuộc phòng nào?" hoặc "X ở phòng nào?" → PHẢI JOIN bảng `phong_ban`.
   - Câu SQL chuẩn:
     + `SELECT nv.ho_ten, pb.ten_phong AS phong_ban FROM nhanvien nv JOIN phong_ban pb ON nv.phong_ban_id = pb.id WHERE nv.ho_ten LIKE '%Tên Nhân Viên%' AND nv.trang_thai_lam_viec != 'Nghỉ việc'`
   - ⚠️ TUYỆT ĐỐI KHÔNG được viết `SELECT phong_ban FROM nhanvien ...` vì cột này không tồn tại trong bảng `nhanvien`.

6. **LUẬT GIAO VIỆC (QUAN TRỌNG - MANY-TO-MANY):**
   - Bảng `cong_viec` KHÔNG lưu trực tiếp người thực hiện (chỉ lưu `nguoi_giao_id`).
   - Để tìm **"Ai làm việc gì"** hoặc **"Việc này ai làm"**:
     => BẮT BUỘC JOIN qua bảng trung gian: `cong_viec_nguoi_nhan`.
   - Lộ trình JOIN chuẩn: `cong_viec` <-> `cong_viec_nguoi_nhan` <-> `nhanvien`.
   - **⚠️ CẢNH BÁO: CÔNG VIỆC CỦA PHÒNG BAN (RẤT QUAN TRỌNG):**
     + **Câu hỏi: "Công việc của phòng kỹ thuật?"** hoặc **"Bao nhiêu công việc phòng kỹ thuật?"**
     + **PHẢI lọc 2 chỗ:**
       1. **Dự án phòng kỹ thuật:** `WHERE d.phong_ban LIKE '%kỹ thuật%'`
       2. **Nhân viên ở phòng kỹ thuật:** `AND nv.phong_ban_id IN (SELECT id FROM phong_ban WHERE ten_phong LIKE '%kỹ thuật%')`
     + **LỲ DO:** Nếu không lọc nhân viên, một nhân viên phòng khác được giao công việc của dự án phòng kỹ thuật sẽ xuất hiện
     + **SQL CHUẨN:**
       ```sql
       SELECT nv.ho_ten, COUNT(cv.id) AS so_cong_viec
       FROM nhanvien nv
       JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id
       JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id
       JOIN du_an d ON cv.du_an_id = d.id
       WHERE d.phong_ban LIKE '%kỹ thuật%'
       AND nv.phong_ban_id IN (SELECT id FROM phong_ban WHERE ten_phong LIKE '%kỹ thuật%')
       GROUP BY nv.id, nv.ho_ten
       ```

22. **LUẬT NĂM VÀO LÀM (RECRUITMENT YEAR LOGIC - RẤT QUAN TRỌNG):**
   - **Câu hỏi: "Nhân viên vào năm 2025?"**
     + Ý nghĩa: Chỉ những người vào LÀM CHÍNH XÁC năm 2025 (không tính những người vào trước đó)
     + SQL: `WHERE YEAR(ngay_vao_lam) = 2025`
   
   - **Câu hỏi: "Vào năm 2025, phòng X có bao nhiêu nhân viên?"** (mơ hồ)
     + Ý nghĩa: Nhân viên hiện tại ở phòng X (có thể vào năm 2025 hoặc trước đó)
     + SQL: `WHERE YEAR(ngay_vao_lam) <= 2025`
   
   - **PHÂN BIỆT RÕ RÀNG:**
     + "vào năm 2025" = `YEAR(ngay_vao_lam) = 2025` (chỉ năm 2025)
     + "từ năm 2025 trở về trước" / "tính đến năm 2025" = `YEAR(ngay_vao_lam) <= 2025`
     + "vào năm 2025 hoặc 2026" = `YEAR(ngay_vao_lam) IN (2025, 2026)`
   
   - **LỰA CHỌN MẶC ĐỊNH (khi câu hỏi mơ hồ):**
     + Nếu câu hỏi chỉ nói "vào năm X" mà không rõ ý -> Dùng `<=` (bao gồm cả những người vào trước)
     + Ví dụ: "2025 công ty tuyển bao nhiêu nhân viên?" -> `= 2025` (chỉ năm đó)
     + Nhưng "2025 phòng X có bao nhiêu nhân viên?" -> `<= 2025` (tất cả nhân viên ở phòng đó từ quá khứ)
7.  **LUẬT CHUẨN HÓA DỮ LIỆU (QUAN TRỌNG - MỚI):**
  - **Trạng thái công việc:** Trong DB lưu chính xác là `'Đã hoàn thành'` (Tuyệt đối không dùng 'Hoàn thành' hay 'Done').
  - **Logic chưa xong:** `trang_thai != 'Đã hoàn thành'`.
  - **Logic trễ hạn:** `han_hoan_thanh <= CURRENT_DATE` AND `trang_thai != 'Đã hoàn thành'`.

8. **LUẬT TRỄ HẠN (DEADLINE LOGIC - RẤT QUAN TRỌNG):**
   - **Định nghĩa:** Một dự án hoặc công việc bị coi là trễ hạn (Overdue) khi:
     `ngay_ket_thuc <= CURRENT_DATE` (hoặc `han_hoan_thanh <= CURRENT_DATE`)
     AND `trang_thai != 'Đã hoàn thành'`.
   - **Lưu ý:** Luôn phải kiểm tra trạng thái. Nếu đã xong (`'Đã hoàn thành'`) thì dù quá ngày cũng không tính là trễ (có thể là xong muộn, nhưng hiện tại không còn treo).
   - **PHÂN BIỆT DỰ ÁN vs CÔNG VIỆC - TRỄ HẠN (RẤT QUAN TRỌNG):**
     + **Câu hỏi: "Dự án trễ hạn?"** (hỏi DỰ ÁN bị trễ)
       → Lọc dự án: `WHERE trang_thai_duan != 'Đã hoàn thành' AND trang_thai_duan != 'Tạm Ngưng' AND ngay_ket_thuc <= CURRENT_DATE`
       → Lý do: Dự án tạm ngưng không tính trễ hạn (đã dừng), chỉ tính những dự án "Đang thực hiện" bị trễ
     + **Câu hỏi: "Công việc trễ hạn?"** (hỏi CÔNG VIỆC bị trễ)
       → Lọc công việc: `WHERE trang_thai != 'Đã hoàn thành' AND han_hoan_thanh <= CURRENT_DATE`
   - **SQL mẫu cho dự án trễ hạn:**
     + `SELECT ten_du_an, ngay_ket_thuc FROM du_an WHERE ngay_ket_thuc <= CURRENT_DATE AND trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm Ngưng')`

23. **LUẬT CHẤM CÔNG & ĐI MUỘN (ATTENDANCE & LATE ARRIVAL - RẤT QUAN TRỌNG):**
   - **Câu hỏi: "Chấm công hôm X của nhân viên Y?"** (Hỏi dữ liệu chấm công bình thường)
     + SQL KHÔNG thêm điều kiện `check_in >= '08:06:00'`
     + Lấy TẤT CẢ dữ liệu: check_in, check_out, bất kể giờ vào là mấy
     + Ví dụ: `SELECT c.ngay, c.check_in, c.check_out FROM cham_cong c WHERE c.ngay = '2026-02-04' AND ...`
   
   - **Câu hỏi: "Ai đi muộn hôm X?"** hoặc **"Danh sách nhân viên đi muộn"** (Hỏi về đi muộn)
     + SQL BẮT BUỘC thêm: `WHERE c.check_in >= '08:06:00'`
     + Chỉ lấy những bản ghi có giờ vào từ 08:06:00 trở đi
     + Ví dụ: `SELECT nv.ho_ten, c.check_in FROM cham_cong c WHERE c.check_in >= '08:06:00' AND ...`
   
   - **PHÂN BIỆT RÕ RÀNG:**
     + Không có từ khóa "muộn", "đi muộn", "late" → Lấy toàn bộ dữ liệu chấm công
     + Có từ khóa "muộn", "đi muộn", "late", "trễ" → Thêm điều kiện `check_in >= '08:06:00'`
   
   - **LUẬT KIỂM SOÁT NGÀY (CRITICAL):**
     + Nếu câu hỏi hỏi về ngày **trong tương lai** (chưa tới) -> CẤM truy vấn chấm công
     + Thay vào đó, hãy trả lời: "Ngày này chưa tới, chưa có dữ liệu chấm công."
     + Ví dụ: Hôm nay là 2026-02-04, hỏi "Ai đi muộn hôm 2026-02-10?" -> Từ chối, không phải trả lời "Không có ai đi muộn"
     + Ngôn ngữ: "Ngày 2026-02-10 chưa tới, tôi không thể cung cấp dữ liệu chấm công cho ngày này."

9. **LUẬT TIẾN ĐỘ & LỊCH SỬ (QUAN TRỌNG NHẤT):**
   - Bảng `cong_viec_tien_do` lưu lịch sử cập nhật (Log). Một việc có nhiều dòng dữ liệu.
   - **Tra cứu đơn lẻ (1 việc):** Dùng `ORDER BY thoi_gian_cap_nhat DESC LIMIT 1` để lấy % mới nhất.
   - **Thống kê/Đếm (Nhiều việc):** BẮT BUỘC dùng Sub-query để lọc ngày mới nhất: 
     `WHERE td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)`.
   - ⛔ **CẤM:** Tuyệt đối KHÔNG dùng `AVG()` hoặc `SUM()` trên cột `phan_tram`.

10. **LUẬT CHI TIẾT QUY TRÌNH (SUB-TASKS):**
   - Khi hỏi về "chi tiết", "các bước", "quy trình" của một việc -> Hãy query bảng `cong_viec_quy_trinh` (lấy cột `ten_buoc`, `trang_thai`).
   - Đừng chỉ lấy mỗi cột `mo_ta` trong bảng `cong_viec` vì nó không đủ chi tiết.
11. **LUẬT TÍNH TIẾN ĐỘ DỰ ÁN (PROJECT PROGRESS RULE):**
   - Bảng `du_an` KHÔNG có cột phần trăm hoàn thành.
   - **Định nghĩa:** Tiến độ dự án = Trung bình cộng (AVG) tiến độ hiện tại của tất cả các công việc (`cong_viec`) thuộc dự án đó.
   - **Công thức SQL bắt buộc:**
     1. Lấy tiến độ mới nhất của từng công việc (dùng Sub-query `MAX(thoi_gian_cap_nhat)`).
     2. Gom nhóm theo dự án (`GROUP BY du_an.id`).
     3. Tính `AVG(phan_tram)`.
     4. Nếu cần lọc (ví dụ > 80%), dùng `HAVING AVG(...) > 80`.
12. **MỐI QUAN HỆ DỰ ÁN - CÔNG VIỆC:**
   - Liên kết: `du_an.id` = `cong_viec.du_an_id`.
   - Tiến độ: `cong_viec.id` = `cong_viec_tien_do.cong_viec_id`
13. **LUẬT TRA CỨU TIẾN ĐỘ AN TOÀN (SAFE JOIN RULE):**
   - Khi tính toán tiến độ dự án hoặc công việc, hãy ưu tiên dùng **`LEFT JOIN cong_viec_tien_do`**.
   - Lý do: Có những dự án mới tạo chưa có log tiến độ. Nếu dùng `INNER JOIN` sẽ bị mất dữ liệu.
   - Xử lý NULL: Sử dụng `COALESCE(AVG(td.phan_tram), 0)` để mặc định là 0% nếu không tìm thấy log.
14. **LUẬT THỐNG KÊ TRẠNG THÁI DỰ ÁN (PROJECT STATUS STATS):**
   - Khi người dùng hỏi thống kê số lượng dự án theo "trạng thái" (VD: Đang thực hiện, Đã xong...):
   - **Không cần tính toán** phức tạp.
   - Truy vấn trực tiếp bảng `du_an`.
   - Sử dụng `GROUP BY trang_thai_duan` (Lưu ý: tên cột là `trang_thai_duan`, KHÔNG dùng `trang_thai` vì đó là cột của bảng công việc).

15. **LUẬT TRA CỨU TIẾN ĐỘ DỰ ÁN (PROJECT PROGRESS - ADVANCED):**
   - **Bối cảnh:** Bảng `du_an` KHÔNG có cột phần trăm.
   - **Logic:** Tiến độ Dự án = Trung bình cộng (AVG) tiến độ *mới nhất* của tất cả công việc (`cong_viec`) thuộc dự án đó.
   - **Công thức SQL BẮT BUỘC (Safe Mode):**
     1. Dùng **`LEFT JOIN`** bảng `cong_viec` và `cong_viec_tien_do` (để không bị mất dự án nếu chưa có log tiến độ).
     2. Xử lý NULL: Dùng `COALESCE(AVG(td.phan_tram), 0)` để mặc định là 0% nếu chưa có dữ liệu.
     3. Lọc mới nhất: `AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)`.
     4. Gom nhóm: `GROUP BY d.id, d.ten_du_an`.

16. **LUẬT DỰ ÁN TẠM NGƯNG (PAUSED PROJECTS):**
    - Khi truy vấn dự án (đặc biệt là dự án Tạm ngưng/Dừng), người dùng luôn muốn biết **Ai chịu trách nhiệm (Leader)**.
    - **Logic lấy tên Leader:** 
      - Bắt buộc JOIN bảng `nhanvien` (alias `nv`).
      - Điều kiện: `du_an.lead_id = nv.id`.
      - Lấy cột: `nv.ho_ten`.
    - **Logic lọc trạng thái:** Dùng `trang_thai LIKE '%Ngưng%'` hoặc `LIKE '%Dừng%'`.
    - **Logic tiến độ:** Vẫn giữ nguyên công thức tính AVG từ bảng `cong_viec` để biết dự án dừng ở mức nào.

17. **LUẬT HIỆU SUẤT NHÂN SỰ (PERFORMANCE):**
    - Đánh giá ai làm việc hiệu quả: Dựa trên số lượng công việc đã hoàn thành (`trang_thai` = 'Đã hoàn thành') và so sánh `ngay_hoan_thanh` <= `han_hoan_thanh` (xong trước hạn).
    - Đánh giá quá tải: Đếm số lượng công việc `trang_thai` = 'Đang thực hiện' của từng người.

18. **LUẬT TÊN CỘT TRẠNG THÁI (STATUS COLUMN NAMES):**
   - LƯU Ý RẤT QUAN TRỌNG VỀ SCHEMA:
     + Bảng `cong_viec` dùng cột: **`trang_thai`** [2].
     + Bảng `du_an` dùng cột: **`trang_thai_duan`** [1].
   - Tuyệt đối không dùng `du_an.trang_thai` (sẽ gây lỗi SQL).

19. **LUẬT DỰ ÁN TẠM NGƯNG:**
    - Khi lọc dự án tạm ngưng, dùng điều kiện: `d.trang_thai_du_an LIKE '%Ngưng%'`.
    - Vẫn tính toán tiến độ trung bình từ `cong_viec` để hiển thị mức độ dở dang.

20. **LUẬT XÁC ĐỊNH CÔNG VIỆC TRỄ HẠN (OVERDUE RULE):**
    - Một công việc bị coi là TRỄ HẠN khi thỏa mãn 2 điều kiện:
      1. `trang_thai` KHÁC 'Đã hoàn thành' (Ví dụ: 'Đang thực hiện', 'Mới tạo'...).
      2. `han_hoan_thanh` < `CURRENT_DATE` (Ngày hiện tại).
    - Câu lệnh SQL mẫu: `WHERE cv.trang_thai != 'Đã hoàn thành' AND cv.han_hoan_thanh < CURDATE()`.

21. **QUY TẮC ĐẾM SỐ LƯỢNG (COUNT RULE) – BẮT BUỘC:**
- KÍCH HOẠT KHI câu hỏi chứa các cụm:
  + "bao nhiêu"
  + "tổng số"
  + "có mấy"
  + "số lượng"
- MỤC TIÊU:
  → Trả lời bằng **SỐ LƯỢNG** (không liệt kê danh sách chi tiết).
- SQL LOGIC BẮT BUỘC:
  → PHẢI sử dụng hàm:
    `COUNT(*) AS total`
- MẪU SQL CHUẨN:
  ```sql
  SELECT COUNT(*) AS total
  FROM <table>;

22. **LUẬT TRA CỨU ĐƠN NGHỈ PHÉP (LEAVE REQUESTS - REAL DATA):**
    - **Cấu trúc bảng `don_nghi_phep` thực tế:**
      + Cột ngày: `ngay_bat_dau` và `ngay_ket_thuc` (KHÔNG dùng `tu_ngay`/`den_ngay`).
      + Khóa ngoại: `nhan_vien_id` (có gạch dưới `_`).
      + Trạng thái: Giá trị lưu là `'da_duyet'` (không dấu, viết thường).
    - **Logic tìm người đang nghỉ:**
      + `CURRENT_DATE` nằm trong khoảng `ngay_bat_dau` và `ngay_ket_thuc`.
      + Điều kiện: `trang_thai = 'da_duyet'`.

23. **LUẬT TRA CỨU QUỸ PHÉP (LEAVE BALANCE - RẤT QUAN TRỌNG):**
    - **Cấu trúc bảng `ngay_phep_nam`:**
      + Khóa ngoại: `nhan_vien_id`.
      + Cột số liệu: `tong_ngay_phep`, `ngay_phep_da_dung`, `ngay_phep_con_lai`.
      + Cột năm: `nam` (lưu năm của kỳ phép).
    - **Logic Join:** `ngay_phep_nam.nhan_vien_id = nhanvien.id`.
    
    - **LUẬT LỌCNĂM (CRITICAL):**
      + Khi hỏi **thông tin phép của nhân viên CỤ THỂ** (bằng id, tên) → BẮT BUỘC lọc năm hiện tại
      + SQL: `WHERE ... AND np.nam = YEAR(CURRENT_DATE)` (năm hiện tại = 2026)
      + Ví dụ: "Phép năm của nhân viên id=5?" → Chỉ trả phép năm 2026
      + Lý do: Bảng ngay_phep_nam có nhiều dòng cho cùng 1 nhân viên (mỗi năm 1 dòng)
      + CẤM truy vấn mà không lọc năm (sẽ return nhiều dòng trùng nhân viên)
    
    - **LUẬT THỐNG KÊ (thống kê toàn công ty/phòng/chức vụ):**
      + Nếu hỏi "tổng phép của toàn công ty" không rõ năm → Lấy năm hiện tại
      + SQL: `SELECT ... WHERE np.nam = YEAR(CURRENT_DATE)` (mặc định)
      + Nếu hỏi rõ "năm 2025" hoặc "tất cả năm" → Không lọc năm (return tất cả)
    
    - **SQL mẫu cho câu hỏi CỤ THỂ:**
      + "Phép năm của Nguyễn Tấn Dũng?" → `SELECT ... WHERE nv.ho_ten LIKE '%Nguyễn Tấn Dũng%' AND np.nam = YEAR(CURRENT_DATE)` (ONLY nv.ho_ten, NOT nv.id)
      + "Phép năm của nhân viên id=5?" → `SELECT ... WHERE nv.id = 5 AND np.nam = YEAR(CURRENT_DATE)` (BEST: Dùng ID thay vì LIKE)
      + **CẢNH BÁO:** Khi dùng LIKE với tên không đầy đủ (VD: LIKE '%Thắng%') có thể match nhiều nhân viên. LƯU Ý: Không được gom/merge dữ liệu từ nhiều nhân viên khác nhau.
    
    - **LOGIC VALIDATION (QUAN TRỌNG):**
      + **Công thức bắt buộc:** `tong_ngay_phep = ngay_phep_da_dung + ngay_phep_con_lai`
      + Khi hiển thị dữ liệu phép năm, PHẢI validate công thức này trước
      + Nếu dữ liệu không thỏa mãn công thức → SẼ CÓ LỖI trong dữ liệu hoặc query, BẮT BUỘC thông báo cho người dùng
      + Ví dụ SAI: "Tổng: 2, Đã dùng: 3, Còn lại: 2" (2 ≠ 3+2) → KHÔNG được trả lời, phải thông báo lỗi

24. **LUẬT TÌM LÃNH ĐẠO / GIÁM ĐỐC (LEADERSHIP LOOKUP):**
    - Khi người dùng hỏi: "Giám đốc là ai?", "Ai là sếp?", "CEO của công ty", "Ban lãnh đạo".
    - **Logic:** Truy vấn bảng `nhanvien`.
    - **Điều kiện:** Tìm kiếm trong cột `chuc_vu` hoặc `vai_tro`.
    - **Từ khóa lọc:** Sử dụng `LIKE '%Giám đốc%'`, `LIKE '%CEO%'`, hoặc `LIKE '%Chủ tịch%'`.
    - **SQL mẫu:** `SELECT ho_ten, chuc_vu, email FROM nhanvien WHERE chuc_vu LIKE '%Giám đốc%' OR chuc_vu LIKE '%CEO%'`.

25. **LUẬT ĐẾM NHÂN VIÊN vs ĐẾM SỰ KIỆN (COUNT vs COUNT DISTINCT - RẤT QUAN TRỌNG):**
    - **PHÂN BIỆT 2 LOẠI CÂU HỎI:**
    
    A. "Bao nhiêu LẦN đi muộn trong tháng 1?"
       + Đếm số **bản ghi chấm công** → `COUNT(*)` hoặc `COUNT(DISTINCT c.id)`
       + Ví dụ: Nếu nhân viên A đi muộn 3 lần → Kết quả = 3
       + SQL: `SELECT COUNT(*) FROM cham_cong c WHERE c.check_in >= '08:06:00' AND MONTH(c.ngay) = 1`
    
    B. "Bao nhiêu NHÂN VIÊN đi muộn trong tháng 1?"
       + Đếm số **nhân viên duy nhất** → `COUNT(DISTINCT c.nhan_vien_id)` hoặc `COUNT(DISTINCT n.id)`
       + Ví dụ: Nếu nhân viên A, B, C mỗi người đi muộn (A 3 lần, B 1 lần, C 2 lần) → Kết quả = 3 (không phải 6)
       + SQL: `SELECT COUNT(DISTINCT c.nhan_vien_id) FROM cham_cong c WHERE c.check_in >= '08:06:00' AND MONTH(c.ngay) = 1`
    
    C. "Tổng nhân viên vắng mặt trong tháng 1?"
       + Đếm số **nhân viên duy nhất** chưa chấm công → `COUNT(DISTINCT id)`
       + SQL: `SELECT COUNT(DISTINCT id) FROM nhanvien WHERE id NOT IN (...) AND trang_thai_lam_viec != 'Nghỉ việc'`
    
    - **QUY TẮC:**
      + Câu hỏi chứa "nhân viên" (số ít/số nhiều) → BẮT BUỘC dùng `COUNT(DISTINCT)`
      + Câu hỏi chứa "lần", "lần", "bao nhiêu lần" → Có thể dùng `COUNT(*)` tùy ngữ cảnh
      + **MẶC ĐỊNH an toàn:** Khi hỏi về nhân viên, luôn dùng `COUNT(DISTINCT nhan_vien_id)`

26. **LUẬT XUẤT FILE BÁO CÁO (REPORT EXPORT - BẮT BUỘC CHI TIẾT):**
    - **Điều kiện kích hoạt:** Khi người dùng yêu cầu xuất file báo cáo (từ khóa: "word", "docx", "file", "báo cáo", "xuất", "in")
    - **QUI TẮC BẮTBUỘC CHO DỮ LIỆU XUẤT FILE:**
      + **PHẢI IN RA TẤT CẢ THÔNG TIN CHI TIẾT**, KHÔNG ĐƯỢC:
        - Tóm tắt dữ liệu
        - Loại bỏ cột
        - Gộp/merge dữ liệu
        - Chỉ hiển thị tóm lược
      + **PHẢI bao gồm:**
        - Tất cả cột từ SQL query
        - Tất cả dòng dữ liệu (không loại bỏ bất kỳ record nào)
        - Định dạng rõ ràng, dễ đọc
      + **VÍ DỤ CHO "BÁOCÁO DỰ ÁN TRỄ HẠN":**
        - ĐÚNG: In danh sách đầy đủ với cột: Tên dự án, Ngày kết thúc, Mô tả, Trạng thái, Người quản lý, v.v.
        - SAI: Chỉ in "Có 3 dự án trễ hạn" (không liệt kê chi tiết)
    - **SCHEMA ĐẦY ĐỦ TRONG FILE:**
      + Tiêu đề báo cáo + ngày tạo
      + Câu hỏi người dùng
      + Kết quả SQL (tất cả dòng + tất cả cột)
      + Không được rút gọn hay tóm lược

27. **BẢNG `cau_hinh_he_thong` (SYSTEM CONFIGURATION):**
    - **Định nghĩa:** Chứa các cấu hình, thiết lập hệ thống, tham số toàn cục
    - **Các cột:** `id`, `ten_cau_hinh` (tên cấu hình), `gia_tri` (giá trị), `mo_ta` (mô tả), `ngay_tao` (ngày tạo)
    - **Câu hỏi liên quan:**
      + "Có bao nhiêu cấu hình hệ thống?" → `SELECT COUNT(*) FROM cau_hinh_he_thong`
      + "Danh sách toàn bộ cấu hình" → `SELECT ten_cau_hinh, gia_tri FROM cau_hinh_he_thong`
      + "Cấu hình nào có tên [X]" → `SELECT * FROM cau_hinh_he_thong WHERE ten_cau_hinh LIKE '%[X]%'`
      + "Cấu hình working_hours được tạo bao giờ?" → `SELECT ngay_tao FROM cau_hinh_he_thong WHERE ten_cau_hinh = 'working_hours'`
    - **Lưu ý:** Đây là bảng riêng biệt, KHÔNG liên quan đến nhân viên, dự án hay công việc
    - **Quyền hạn:** Thường là thông tin dành cho quản trị viên hệ thống

28. **BẢNG `cong_viec_quy_trinh` (WORKFLOW STEPS - QUAN TRỌNG):**
    - **Định nghĩa:** Mỗi công việc (cong_viec) có thể được chia thành nhiều bước (quy trình)
    - **Các cột:** `id`, `cong_viec_id` (FK), `ten_buoc` (tên bước), `mo_ta` (mô tả), `trang_thai` (trạng thái), `ngay_bat_dau`, `ngay_ket_thuc`, `ngay_tao`
    - **Trạng thái bước:** "Đang thực hiện", "Đã hoàn thành", "Hủy", "Tạm dừng", v.v.
    - **⚠️ QUY TẮC NHẤT QUÁN GIỮA COUNT VÀ SELECT (CRITICAL):**
      + **PHẢI dùng cùng điều kiện WHERE cho COUNT và SELECT**
      + **KHÔNG được:** COUNT lấy tất cả 10 bước, nhưng SELECT chỉ lấy 8 bước (vì có thêm filter trạng thái)
      + **ĐÚNG:** Nếu lọc bước "Đã hoàn thành", cả COUNT và SELECT phải có `WHERE ... AND trang_thai = 'Đã hoàn thành'`
    - **QUY TẮC LỌC MẶC ĐỊNH:**
      + Khi hỏi về bước, MẶC ĐỊNH chỉ lấy bước HOẠT ĐỘNG (không lấy bước bị "Hủy")
      + Điều kiện: `WHERE cong_viec_id = ? AND trang_thai != 'Hủy'` (hoặc phù hợp với DB)
      + LỆ NGOẠI: Nếu người dùng tường minh hỏi "Các bước bị hủy", thì mới bỏ filter
    - **SQL examples:**
      + "Công việc 205 có bao nhiêu bước?" → `SELECT COUNT(*) FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND trang_thai != 'Hủy'`
      + "Liệt kê bước của công việc 205" → `SELECT ten_buoc, trang_thai FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND trang_thai != 'Hủy' ORDER BY ngay_bat_dau`
      + CHÍNH XÁC: COUNT = 8, SELECT = 8 bước (KHÔNG bao gồm bước bị hủy)

29. **SUBQUERY & IN OPERATOR - QUAN TRỌNG ĐỀ PHÒNG LỖI:**
    - **Vấn đề:** Khi query công việc theo tên (hoặc bất kỳ điều kiện nào), có thể có NHIỀU công việc khớp
    - **TUYỆT ĐỐI CẤM:** Dùng `WHERE cong_viec_id = (SELECT id FROM cong_viec WHERE ...)` khi subquery có thể trả >1 row
    - **PHẢI dùng:** `WHERE cong_viec_id IN (SELECT id FROM cong_viec WHERE ...)`
    - **Ví dụ SAI (lỗi):**
      + `SELECT COUNT(*) FROM cong_viec_quy_trinh WHERE cong_viec_id = (SELECT id FROM cong_viec WHERE ten_cong_viec LIKE '%Oracle Cloud%')`
      + Nếu có 2 công việc "Oracle Cloud", subquery trả 2 row → ERROR: "Subquery returns more than 1 row"
    - **Ví dụ ĐÚNG:**
      + `SELECT COUNT(*) FROM cong_viec_quy_trinh WHERE cong_viec_id IN (SELECT id FROM cong_viec WHERE ten_cong_viec LIKE '%Oracle Cloud%')`
      + Kết quả: Tính tổng bước của TẤT CẢ công việc có tên chứa "Oracle Cloud"
    - **Nguyên tắc:**
      + ALWAYS dùng `IN` cho subquery nhiều hàng
      + Chỉ dùng `=` khi chắc chắn subquery trả 1 hàng duy nhất (VD: `WHERE id = (SELECT ... LIMIT 1)`)

30. **BẢNG `lich_trinh` (SCHEDULE/CALENDAR - LỊ TRÌNH SỰ KIỆN):**
    - **Định nghĩa:** Lưu trữ lịch trình, sự kiện, hội thảo, buổi gặp mặt, v.v.
    - **Các cột:** `id`, `tieu_de` (tiêu đề sự kiện), `mo_ta` (mô tả chi tiết), `ngay_bat_dau`, `ngay_ket_thuc`, `ngay_tao`
    - **Câu hỏi liên quan:**
      + "Có bao nhiêu sự kiện?" → `SELECT COUNT(*) FROM lich_trinh`
      + "Sự kiện nào hôm nay?" → `SELECT tieu_de, mo_ta FROM lich_trinh WHERE ngay_bat_dau = CURRENT_DATE`
      + "Sự kiện sắp tới?" → `SELECT tieu_de, ngay_bat_dau FROM lich_trinh WHERE ngay_bat_dau >= CURRENT_DATE ORDER BY ngay_bat_dau`
      + "Sự kiện nào đã qua?" → `SELECT tieu_de, ngay_ket_thuc FROM lich_trinh WHERE ngay_ket_thuc < CURRENT_DATE ORDER BY ngay_ket_thuc DESC`
      + "Tìm sự kiện 'LBS startup'?" → `SELECT tieu_de, mo_ta, ngay_bat_dau FROM lich_trinh WHERE tieu_de LIKE '%LBS%'`
    - **Lưu ý:** LUÔN kiểm tra `ngay_bat_dau >= CURRENT_DATE` (sắp tới) vs `<` (đã qua)

31. **BẢNG `quyen` (PERMISSIONS - QUYỀN HẠN VÀ PHÂN QUYỀN):**
    - **Định nghĩa:** Lưu trữ danh sách các quyền hạn, nhóm quyền, và mã quyền trong hệ thống.
    - **Các cột:** `id`, `nhom_quyen` (nhóm quyền: phongban, nhanvien, quanly, admin), `ma_quyen` (mã quyền duy nhất), `ten_quyen` (tên hiển thị của quyền)
    - **Câu hỏi liên quan:**
      + "Có bao nhiêu quyền hạn?" → `SELECT COUNT(*) FROM quyen`
      + "Liệt kê tất cả quyền hạn" → `SELECT ma_quyen, ten_quyen, nhom_quyen FROM quyen`
      + "Quyền hạn nào liên quan tới phòng ban?" → `SELECT ma_quyen, ten_quyen FROM quyen WHERE nhom_quyen LIKE '%phongban%'`
      + "Tìm quyền 'Thêm phòng ban'" → `SELECT id, ma_quyen, ten_quyen FROM quyen WHERE ten_quyen LIKE '%Thêm phòng ban%'`
      + "Danh sách quyền hạn của nhóm 'nhanvien'" → `SELECT ma_quyen, ten_quyen FROM quyen WHERE nhom_quyen = 'nhanvien'`
    - **Lưu ý:** `nhom_quyen` là nhóm lôgic (phongban, nhanvien, quanly, admin); `ma_quyen` là mã duy nhất không thay đổi

32. **BẢNG `tai_lieu` (DOCUMENTS - TÀI LIỆU & HỆ THỐNG):**
    - **Định nghĩa:** Lưu trữ tài liệu, báo cáo, hộp cò và các file chia sẻ trong hệ thống.
    - **Các cột:** `id`, `ten_tai_lieu` (tên tài liệu), `mo_ta` (mô tả), `file_path` (đường dẫn file), `file_name` (tên file), `file_type` (MIME type), `file_size` (kích thước), `loai_tai_lieu` (loại: Báo cáo, Hộp cò, v.v.), `trang_thai` (Hoạt động, Đã xóa), `doi_tuong_xem` (Tất cả, Phòng ban, Nhân viên), `luot_xem` (số lần xem), `luot_tai` (số lần tải), `ngay_tao`, `ngay_cap_nhat`, `nguoi_tao_id`
    - **Câu hỏi liên quan:**
      + "Có bao nhiêu tài liệu?" → `SELECT COUNT(*) FROM tai_lieu WHERE trang_thai != 'Đã xóa'`
      + "Liệt kê tất cả tài liệu" → `SELECT ten_tai_lieu, loai_tai_lieu, ngay_tao FROM tai_lieu WHERE trang_thai != 'Đã xóa'`
      + "Tài liệu nào được xem nhiều nhất?" → `SELECT ten_tai_lieu, luot_xem FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY luot_xem DESC LIMIT 1`
      + "Tài liệu nào được tải nhiều nhất?" → `SELECT ten_tai_lieu, loai_tai_lieu, luot_tai FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY luot_tai DESC LIMIT 1`
      + "Tìm tài liệu 'File cài đặt CSA'" → `SELECT ten_tai_lieu, file_name, file_size, ngay_tao FROM tai_lieu WHERE ten_tai_lieu LIKE '%File cài đặt CSA%'`
    - **Lưu ý:** Luôn lọc `WHERE trang_thai != 'Đã xóa'` để không hiển thị file đã bị xóa; dùng `luot_xem` và `luot_tai` để sắp xếp theo mức độ sử dụng

SCHEMA CHI TIẾT:
{HRM_SCHEMA_RAW}
"""

# ==========================================================
import pandas as pd
import re
from langchain_core.prompts import PromptTemplate


# ==========================================================
# ROUTER: PHÂN LOẠI CÂU HỎI HRM_SQL vs STATIC_DOC
# ==========================================================
ROUTER_PROMPT = ChatPromptTemplate.from_template("""
Bạn là bộ phân loại câu hỏi cho chatbot ICS.

NHIỆM VỤ: Với mỗi câu hỏi, hãy trả lời duy nhất một từ trong 2 từ sau đây:
- HRM_SQL  -> nếu câu hỏi thuộc về quản lý nhân sự, chấm công, lương, dự án, công việc, phép năm, hệ thống HRM nội bộ.
- STATIC_DOC -> nếu câu hỏi thuộc về giới thiệu công ty ICS, sản phẩm, giải pháp, tiêu chuẩn, quy trình, ISO, khách hàng, case study, ví dụ triển khai, marketing.

Quy tắc:
- Nếu câu hỏi hỏi về nhân viên, phòng ban, dự án nội bộ, công việc, chấm công, phép năm, hệ thống HRM, lịch trình/sự kiện nội bộ (bảng `lich_trinh`) → chọn HRM_SQL.
- Các câu hỏi có chứa từ "sự kiện", "su kien", "lịch trình", "lich trinh", "lịch họp", "lich hop" **luôn** thuộc phạm vi HRM_SQL (tra cứu lịch/sự kiện nội bộ), KHÔNG được chọn STATIC_DOC.
- Nếu câu hỏi hỏi "ICS là ai", "ICS làm gì", sản phẩm/giải pháp ICS, tiêu chuẩn ISO 27001, VietGuard, Smart Dashboard, khách hàng, dự án đã triển khai cho khách → chọn STATIC_DOC.
- Nếu mơ hồ nhưng có thể map được sang HRM → ƯU TIÊN HRM_SQL.

CÂU HỎI: {question}

Chỉ trả lời đúng MỘT từ: HRM_SQL hoặc STATIC_DOC.
""")

router_chain = ROUTER_PROMPT | llm_router | StrOutputParser()
# Nhớ import các hàm tạo file chúng ta đã viết ở bước trước
# from report_generator import create_word_report, create_pdf_report (hoặc để chung file cũng được)

# --- 1. HÀM SINH SQL TỪ LLM ---
def generate_sql_from_llm(question):
    """
    Gửi Schema và câu hỏi cho AI để nhận lại câu lệnh SQL
    """
    template = f"""
    {HRM_SCHEMA_ENHANCED}
    
    Dựa trên quy tắc và schema trên, hãy viết câu lệnh SQL để trả lời câu hỏi: "{question}"
    
    Yêu cầu:
    - Chỉ trả về duy nhất câu lệnh SQL. 
    - Không giải thích, không markdown (```sql).
    - Nếu cần xuất file, hãy lấy càng nhiều cột chi tiết càng tốt.
    """
    
    # Giả sử bạn đã khởi tạo biến 'llm' (OpenAI/Google Gemini) ở đầu file
    # response = llm.invoke(template) 
    # return response.content.strip().replace("```sql", "").replace("```", "")
    
    # [CODE MẪU CHO LANGCHAIN]:
    prompt = PromptTemplate.from_template(template)
    chain = prompt | llm 
    sql = chain.invoke({})
    
    # Làm sạch chuỗi SQL (xóa markdown thừa nếu có)
    sql_clean = sql.strip().replace("```sql", "").replace("```", "").strip()
    return sql_clean

# --- 2. HÀM TÓM TẮT KẾT QUẢ (NÓI CHUYỆN VỚI SẾP) ---
def generate_natural_response(question, data):
    """
    AI đọc dữ liệu SQL và trả lời Sếp bằng tiếng Việt tự nhiên
    """
    if not data:
        return "Thưa sếp, em đã tìm trong hệ thống nhưng không thấy dữ liệu nào phù hợp ạ."
        
    data_preview = str(data[:10]) # Chỉ đưa 10 dòng đầu cho AI đọc để tiết kiệm token
    
    prompt = f"""
    Câu hỏi của Sếp: "{question}"
    Dữ liệu tìm được từ Database: {data_preview}
    
    Hãy đóng vai trợ lý ảo chuyên nghiệp, trả lời ngắn gọn, đi vào trọng tâm.
    Nếu dữ liệu là danh sách dài, hãy chỉ tóm tắt các con số quan trọng (Tổng số, Top đầu...).
    """
    
    return llm.invoke(prompt).content

# --- 3. HÀM XỬ LÝ CHÍNH (MAIN HANDLER) ---
def handle_query(question):
    """
    Hàm này sẽ được ui.py gọi.
    Input: Câu hỏi của user.
    Output: Dictionary chứa nội dung trả lời và thông tin file (nếu có).
    """
    print(f"DEBUG: Nhận câu hỏi: {question}")
    
    try:
        # BƯỚC 1: AI Dịch câu hỏi sang SQL
        sql_query = generate_sql_from_llm(question)
        print(f"DEBUG: SQL Generated: {sql_query}")
        
        # BƯỚC 2: Chạy SQL lấy dữ liệu thô
        # (Giả sử bạn đã có hàm execute_sql_query kết nối DB)
        raw_data = execute_sql_query(sql_query) 
        
        # Nếu không có dữ liệu hoặc lỗi
        if isinstance(raw_data, str) and "Error" in raw_data:
            return {
                "type": "text", 
                "content": f"Hệ thống gặp lỗi khi truy vấn: {raw_data}"
            }
        
        if not raw_data:
            return {
                "type": "text", 
                "content": "Dạ em kiểm tra thì không thấy dữ liệu nào khớp với yêu cầu của Sếp ạ."
            }

        # BƯỚC 3: PHÂN TÍCH Ý ĐỊNH XUẤT FILE
        # Kiểm tra xem Sếp có đòi file không
        q_lower = question.lower()
        export_needed = False
        file_path = None
        file_format = None
        
        if "word" in q_lower or "docx" in q_lower or "văn bản" in q_lower:
            export_needed = True
            file_format = "docx"
            # Gọi hàm tạo Word (đã viết ở bước trước)
            file_path = create_word_report(raw_data, title="BÁO CÁO HRM", filename_prefix="baocao")
            
        elif "pdf" in q_lower or "xuất file" in q_lower: # Mặc định xuất PDF nếu nói chung chung
            export_needed = True
            file_format = "pdf"
            # Gọi hàm tạo PDF
            file_path = create_pdf_report(raw_data, title="BAO CAO HRM", filename_prefix="baocao")

        # BƯỚC 4: TRẢ KẾT QUẢ VỀ UI
        if export_needed and file_path:
            return {
                "type": "file",
                "content": f"Dạ, em đã trích xuất xong dữ liệu Sếp cần ({len(raw_data)} dòng). Mời Sếp tải báo cáo bên dưới ạ:",
                "path": file_path,
                "format": file_format
            }
        else:
            # Nếu không xuất file, nhờ AI tóm tắt bằng lời
            summary = generate_natural_response(question, raw_data)
            return {
                "type": "text",
                "content": summary
            }

    except Exception as e:
        print(f"ERROR: {str(e)}")
        return {"type": "text", "content": "Xin lỗi Sếp, hệ thống đang gặp chút trục trặc kỹ thuật."}
# ==========================================================

# --- PROMPT 1: SINH SQL (Kèm Few-Shot Learning) ---
SQL_PROMPT = ChatPromptTemplate.from_template("""
Bạn là SQL Generation Engine. Nhiệm vụ: Chuyển câu hỏi thành SQL Server/MySQL query tối ưu.

⛔ BỘ LUẬT CẤM (CRITICAL RULES):
1. **Output:** Chỉ trả về code SQL trần (Raw text). KHÔNG Markdown, KHÔNG giải thích.
2. **Luật Đi Muộn:** Bắt buộc `check_in >= '08:06:00'`.
3. **Luật Vắng Mặt:** Dùng `NOT IN (SELECT...)`.
4. **An toàn:** Chỉ dùng bảng/cột có trong SCHEMA.
5. **⚠️ LIỆT KÊ = SELECT TẤT CẢ CỘT (CRITICAL):**
   - Khi câu hỏi có "liệt kê", "danh sách", "cho tôi xem" → PHẢI SELECT TẤT CẢ CỘT LIÊN QUAN
   - KHÔNG được SELECT ít cột (vd: chỉ SELECT tên mà bỏ ngày/trạng thái)
   - KHÔNG được DROP cột, vì AI cần tất cả info để liệt kê chi tiết
   - Ví dụ ĐÚNG: "Liệt kê dự án trễ hạn" → SELECT ten_du_an, ngay_ket_thuc, mo_ta, trang_thai_duan (tất cả cột)
   - Ví dụ SAI: "Liệt kê dự án trễ hạn" → SELECT ten_du_an (thiếu ngày, trạng thái)
6. Ngoài lề:
- Chỉ trả về "NO_DATA" nếu:
  a) Câu hỏi hoàn toàn KHÔNG liên quan đến HRM / Dự án / Nhân sự / Sự kiện nội bộ
  b) Không ánh xạ được tới BẤT KỲ bảng nào trong schema
- Tuyệt đối KHÔNG trả về "NO_DATA" cho các câu hỏi có chứa từ khóa "sự kiện", "su kien", "lịch trình", "lich trinh", "lịch họp", "lich hop" – đây là các câu hỏi thuộc phạm vi HRM, PHẢI dùng bảng `lich_trinh`.
- Nếu câu hỏi còn mơ hồ nhưng có khả năng liên quan, hãy suy luận hợp lý nhất và sinh SQL an toàn.

7. **⛔ LỌC KPI (CONFIDENTIAL - KHÔNG CÔNG KHAI):**
   - Nếu câu hỏi chứa từ khóa: "", "performance", "target", "kết quả kinh doanh"
   - Phải TỪNG VÀ KHÔNG SINH SQL
   - Trả về: "KPI_BLOCKED" (để backend xử lý thành: "Dữ liệu KPI là thông tin không công khai. Bạn có muốn hỏi về các vấn đề khác không?")
   - KHÔNG bao giờ cho phép query đó dù câu hỏi có hint rằng dữ liệu ở đâu

8. **⛔ LỌC LỊCH SỬ NHÂN SỰ (CONFIDENTIAL - KHÔNG CÔNG KHAI):**
   - Nếu câu hỏi chứa từ khóa: "lịch sử", "quá khứ", "công việc trước", "kinh nghiệm làm việc", "hồ sơ", "tiểu sử","từng làm", "đã làm việc ở"
   - Phải TỪNG VÀ KHÔNG SINH SQL
   - Trả về: "HISTORY_BLOCKED" (để backend xử lý thành: "Dữ liệu lịch sử nhân sự là thông tin không được công khai. Bạn có muốn hỏi về các vấn đề khác không?")
   - KHÔNG bao giờ cho phép query đó dù câu hỏi có hint rằng dữ liệu ở đâu

HỌC TỪ VÍ DỤ (FEW-SHOT):

⚠️ **CẢNH BÁO: PHÂN BIỆT LIỆT KÊ (SELECT) vs ĐẾM (COUNT)**
- Câu có từ "Liệt kê", "Danh sách", "Cho tôi xem" → Dùng SELECT chi tiết, ORDER BY
- Câu có từ "Bao nhiêu", "Tổng", "Đếm", "Thống kê" → Dùng COUNT, GROUP BY
- Ví dụ:
  + "Liệt kê dự án đang thực hiện theo phòng" → SELECT phong_ban, ten_du_an ... ORDER BY
  + "Bao nhiêu dự án đang thực hiện ở từng phòng" → SELECT phong_ban, COUNT(*) ... GROUP BY

- User: "Hôm nay ai đi muộn?" 
  -> SQL: SELECT n.ho_ten, c.check_in FROM cham_cong c JOIN nhanvien n ON c.nhan_vien_id = n.id WHERE c.ngay = CURRENT_DATE AND c.check_in >= '08:06:00'

- User: "Ai vắng mặt hôm nay?"
  -> SQL: SELECT ho_ten FROM nhanvien WHERE id NOT IN (SELECT nhan_vien_id FROM cham_cong WHERE ngay = CURRENT_DATE)

User: "Lương cơ bản của Nam là bao nhiêu?"
  -> SQL: SELECT ho_ten, luong_co_ban FROM nhanvien WHERE ho_ten LIKE '%Nam%'
                                              
- User: "Có dự án nào đang bị trễ hạn không?"
  -> SQL: SELECT ten_du_an, ngay_ket_thuc FROM du_an WHERE ngay_ket_thuc < CURRENT_DATE AND trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm Ngưng')

- User: "Liệt kê các dự án quá hạn và tên người quản lý?"
  -> SQL: SELECT d.ten_du_an, n.ho_ten, d.ngay_ket_thuc FROM du_an d JOIN nhanvien n ON d.lead_id = n.id WHERE d.ngay_ket_thuc < CURRENT_DATE AND d.trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm Ngưng')

- User: "Tiến độ hiện tại của công việc 'Lên phương án hợp tác với TPX' đến đâu rồi?"
  -> SQL: SELECT td.phan_tram, td.thoi_gian_cap_nhat FROM cong_viec_tien_do td JOIN cong_viec cv ON td.cong_viec_id = cv.id WHERE cv.ten_cong_viec LIKE '%Lên phương án hợp tác với TPX%' ORDER BY td.thoi_gian_cap_nhat DESC LIMIT 1

- User: "Cho tôi xem chi tiết các bước của việc 'Làm việc với a Bình BIDV'?"
  -> SQL: SELECT qt.ten_buoc, qt.trang_thai, qt.mo_ta, qt.ngay_ket_thuc FROM cong_viec_quy_trinh qt JOIN cong_viec cv ON qt.cong_viec_id = cv.id WHERE cv.ten_cong_viec LIKE '%Tuyển dụng nhân sự%' ORDER BY qt.ngay_bat_dau ASC

- User: "Công việc nào đang thực hiện của phòng kỹ thuật?"
  -> SQL: SELECT DISTINCT cv.ten_cong_viec FROM cong_viec cv JOIN du_an d ON cv.du_an_id = d.id WHERE cv.trang_thai = 'Đang thực hiện' AND d.phong_ban LIKE '%kỹ thuật%'

- User: "Công việc nào chưa xong của phòng bán hàng?"
  -> SQL: SELECT DISTINCT cv.ten_cong_viec, cv.han_hoan_thanh FROM cong_viec cv JOIN du_an d ON cv.du_an_id = d.id WHERE cv.trang_thai != 'Đã hoàn thành' AND d.phong_ban LIKE '%bán hàng%'

- User: "Bao nhiêu công việc đang thực hiện ở từng phòng ban?"
  -> SQL: SELECT d.phong_ban, COUNT(DISTINCT cv.id) AS so_cong_viec FROM cong_viec cv JOIN du_an d ON cv.du_an_id = d.id WHERE cv.trang_thai = 'Đang thực hiện' GROUP BY d.phong_ban

# --- ⚠️ CRITICAL: CÔNG VIỆC CỦA PHÒNG BAN PHẢI LỌC NHÂN VIÊN CÙNG PHÒNG ---
- User: "Bao nhiêu công việc nhân viên phòng kỹ thuật nhận?"
  -> SQL: SELECT nv.ho_ten, COUNT(cv.id) AS so_cong_viec FROM nhanvien nv JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id JOIN du_an d ON cv.du_an_id = d.id WHERE d.phong_ban LIKE '%kỹ thuật%' AND nv.phong_ban_id IN (SELECT id FROM phong_ban WHERE ten_phong LIKE '%kỹ thuật%') GROUP BY nv.id, nv.ho_ten

- User: "Liệt kê công việc của nhân viên phòng kỹ thuật"
  -> SQL: SELECT nv.ho_ten, cv.ten_cong_viec, cv.han_hoan_thanh, cv.trang_thai FROM nhanvien nv JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id JOIN du_an d ON cv.du_an_id = d.id WHERE d.phong_ban LIKE '%kỹ thuật%' AND nv.phong_ban_id IN (SELECT id FROM phong_ban WHERE ten_phong LIKE '%kỹ thuật%') ORDER BY nv.ho_ten

User: "Liệt kê các công việc đã hoàn thành trên 50%?"
  -> SQL: SELECT cv.ten_cong_viec, td.phan_tram, td.thoi_gian_cap_nhat FROM cong_viec cv JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id WHERE td.phan_tram > 50 AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
                                              
- User: "Có bao nhiêu công việc đã hoàn thành trên 50%?"
  -> SQL: SELECT COUNT(cv.id) AS so_luong FROM cong_viec cv JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id WHERE td.phan_tram > 50 AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)                        

- User: "Dự án nào được bắt đầu gần đây nhất?"
  -> SQL: SELECT ten_du_an, ngay_bat_dau FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng' ORDER BY ngay_bat_dau DESC LIMIT 1

- User: "Dự án nào kết thúc gần đây nhất?"
  -> SQL: SELECT ten_du_an, ngay_ket_thuc FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng' ORDER BY ngay_ket_thuc DESC LIMIT 1
- User: "Dự án nào bắt đầu lâu nhất?"
  -> SQL: SELECT ten_du_an, ngay_bat_dau FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng' ORDER BY ngay_bat_dau ASC LIMIT 1

- User: "Dự án nào kết thúc lâu nhất?"
  -> SQL: SELECT ten_du_an, ngay_ket_thuc FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng' ORDER BY ngay_ket_thuc ASC LIMIT 1

- User: "Có bao nhiêu dự án?"
  -> SQL: SELECT COUNT(*) as so_luong FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng'

- User: "Danh sách tất cả dự án"
  -> SQL: SELECT ten_du_an, trang_thai_duan FROM du_an WHERE trang_thai_duan != 'Tạm Ngưng'

- User: "Danh sách tất cả dự án kể cả bị tạm dừng"
  -> SQL: SELECT ten_du_an, trang_thai_duan FROM du_an

- User: "Dự án nào đang thực hiện?"
  -> SQL: SELECT ten_du_an FROM du_an WHERE trang_thai_duan = 'Đang thực hiện'

- User: "Dự án nào bị tạm dừng?"
  -> SQL: SELECT ten_du_an FROM du_an WHERE trang_thai_duan = 'Tạm Ngưng'

- User: "Dự án của phòng Marketing?"
  -> SQL: SELECT phong_ban, ten_du_an FROM du_an WHERE phong_ban LIKE '%Marketing%' AND trang_thai_duan != 'Tạm Ngưng'

- User: "Dự án đang thực hiện của phòng Marketing?"
  -> SQL: SELECT phong_ban, ten_du_an FROM du_an WHERE phong_ban LIKE '%Marketing%' AND trang_thai_duan = 'Đang thực hiện'

- User: "Trạng thái đang chạy theo từng phòng ban?"
  -> SQL: SELECT phong_ban, ten_du_an FROM du_an WHERE trang_thai_duan = 'Đang thực hiện' ORDER BY phong_ban

- User: "Bao nhiêu dự án đang thực hiện ở từng phòng ban?"
  -> SQL: SELECT phong_ban, COUNT(id) as so_du_an FROM du_an WHERE trang_thai_duan = 'Đang thực hiện' GROUP BY phong_ban

- User: "Ai quản lý dự án 'Web HRM'?"
  -> SQL: SELECT nv.ho_ten FROM du_an d JOIN nhanvien nv ON d.lead_id = nv.id WHERE d.ten_du_an LIKE '%Web HRM%'

User: "Thống kê số lượng dự án theo từng trạng thái?"
  -> SQL: SELECT trang_thai_du_an, COUNT(id) FROM du_an GROUP BY trang_thai_du_an
                                              
User: "Liệt kê những dự án đã hoàn thành trên 80%?"
  -> SQL: SELECT d.ten_du_an, AVG(td.phan_tram) as tien_do_tb FROM du_an d JOIN cong_viec cv ON d.id = cv.du_an_id JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id WHERE td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id) AND d.trang_thai_duan != 'Tạm Ngưng' GROUP BY d.id, d.ten_du_an HAVING AVG(td.phan_tram) > 80          

 User: "Có bao nhiêu dự án có tiến độ dưới 50%?"
  -> SQL: SELECT COUNT(*) as so_luong FROM (SELECT d.id FROM du_an d JOIN cong_viec cv ON d.id = cv.du_an_id JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id WHERE td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id) AND d.trang_thai_duan != 'Tạm Ngưng' GROUP BY d.id HAVING AVG(td.phan_tram) < 50) as subquery

- User: "Liệt kê các dự án có tiến độ dưới 50%?"
  -> SQL: SELECT d.ten_du_an, AVG(td.phan_tram) as tien_do_trung_binh FROM du_an d JOIN cong_viec cv ON d.id = cv.du_an_id JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id WHERE td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id) AND d.trang_thai_duan != 'Tạm Ngưng' GROUP BY d.id, d.ten_du_an HAVING AVG(td.phan_tram) < 50                                              
                                              
     

- User: "Tiến độ dự án 'Database Mobifone' hiện tại là bao nhiêu?"
  -> SQL: SELECT d.ten_du_an, COALESCE(AVG(td.phan_tram), 0) as phan_tram_hoan_thanh 
          FROM du_an d 
          LEFT JOIN cong_viec cv ON d.id = cv.du_an_id 
          LEFT JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id 
          AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
          WHERE d.ten_du_an LIKE '%Database Mobifone%'
          GROUP BY d.id, d.ten_du_an                                            

- User: "Thống kê số lượng dự án theo từng trạng thái?"
  -> SQL: SELECT trang_thai_duan, COUNT(id) as so_luong FROM du_an GROUP BY trang_thai_duan

- User: "Có bao nhiêu dự án đang ở trạng thái 'Đang thực hiện'?"
  -> SQL: SELECT COUNT(id) as so_luong FROM du_an WHERE trang_thai_du_an LIKE '%Đang thực hiện%'                                                                                          

- User: "Những dự án nào đang bị tạm ngưng và ai là quản lý?"
  -> SQL: SELECT d.ten_du_an, d.trang_thai, COALESCE(AVG(td.phan_tram), 0) as tien_do_luc_dung, nv.ho_ten as quan_ly_du_an
          FROM du_an d 
          LEFT JOIN cong_viec cv ON d.id = cv.du_an_id 
          LEFT JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id 
          AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
          LEFT JOIN nhanvien nv ON d.lead_id = nv.id
          WHERE d.trang_thai LIKE '%Ngưng%' OR d.trang_thai LIKE '%Dừng%'
          GROUP BY d.id, d.ten_du_an, d.trang_thai, nv.ho_ten

# --- Kịch bản: Hỏi thông tin Lead của một dự án cụ thể ---
- User: "Ai đang phụ trách dự án 'Oracle Cloud' và tiến độ thế nào?"
  -> SQL: SELECT d.ten_du_an, nv.ho_ten as lead_du_an, nv.email, COALESCE(AVG(td.phan_tram), 0) as tien_do
          FROM du_an d 
          LEFT JOIN nhanvien nv ON d.lead_id = nv.id
          LEFT JOIN cong_viec cv ON d.id = cv.du_an_id 
          LEFT JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id 
          AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
          WHERE d.ten_du_an LIKE '%Oracle Cloud%'
          GROUP BY d.id, d.ten_du_an, nv.ho_ten, nv.email   

- User: "Top 5 nhân viên hoàn thành nhiều công việc nhất trong tháng này?"
  -> SQL: SELECT nv.ho_ten, COUNT(cv.id) as so_viec_hoan_thanh, pb.ten_phong
          FROM nhanvien nv 
          JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id 
          JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id 
          JOIN phong_ban pb ON nv.phong_ban_id = pb.id
          WHERE cv.trang_thai = 'Đã hoàn thành' AND MONTH(cv.ngay_hoan_thanh) = MONTH(CURRENT_DATE())
          GROUP BY nv.id, nv.ho_ten, pb.ten_phong
          ORDER BY so_viec_hoan_thanh DESC LIMIT 5

- User: "Thống kê khối lượng công việc đang chạy theo từng phòng ban?"
  -> SQL: SELECT pb.ten_phong, COUNT(cv.id) as so_luong_viec_dang_lam 
          FROM phong_ban pb 
          JOIN cong_viec cv ON pb.id = cv.phong_ban_id 
          WHERE cv.trang_thai = 'Đang thực hiện' 
          GROUP BY pb.ten_phong 
          ORDER BY so_luong_viec_dang_lam DESC

- User: "Những dự án nào đang bị tạm ngưng và ai là quản lý?"
  -> SQL: SELECT d.ten_du_an, d.trang_thai_duan, COALESCE(AVG(td.phan_tram), 0) as tien_do_luc_dung, nv.ho_ten as quan_ly_du_an
          FROM du_an d 
          LEFT JOIN cong_viec cv ON d.id = cv.du_an_id 
          LEFT JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id 
          AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
          LEFT JOIN nhanvien nv ON d.lead_id = nv.id
          WHERE d.trang_thai_duan LIKE '%Ngưng%' OR d.trang_thai_duan LIKE '%Dừng%'
          GROUP BY d.id, d.ten_du_an, d.trang_thai_duan, nv.ho_ten

- User: "Thống kê số lượng dự án theo từng trạng thái?"
  -> SQL: SELECT trang_thai_duan, COUNT(id) as so_luong FROM du_an GROUP BY trang_thai_duan                                              

- User: "Kiểm tra xem Trần Đình Nam có công việc nào đang bị trễ hạn không?"
  -> SQL: SELECT cv.ten_cong_viec, cv.han_hoan_thanh, cv.trang_thai, nv.ho_ten
          FROM cong_viec cv
          JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
          JOIN nhanvien nv ON cvnn.nhan_vien_id = nv.id
          WHERE nv.ho_ten LIKE '%Trần Đình Nam%'
          AND cv.trang_thai != 'Đã hoàn thành' 
          AND cv.han_hoan_thanh < CURRENT_DATE


- User: "Liệt kê các công việc đã làm xong của nhân viên mã số 24?"
  -> SQL: SELECT cv.ten_cong_viec, cv.ngay_hoan_thanh, cv.muc_do_uu_tien
          FROM cong_viec cv
          JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
          WHERE cvnn.nhan_vien_id = 24
          AND cv.trang_thai = 'Đã hoàn thành'


- User: "Danh sách công việc và tình trạng hạn chót của dự án Web HRM?"
  -> SQL: SELECT cv.ten_cong_viec, nv.ho_ten as nguoi_lam, cv.han_hoan_thanh, cv.trang_thai,
                 CASE 
                    WHEN cv.trang_thai != 'Đã hoàn thành' AND cv.han_hoan_thanh < CURRENT_DATE THEN 'Trễ hạn'
                    ELSE 'Đúng hạn/Đang chạy'
                 END as tinh_trang_han
          FROM cong_viec cv
          JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
          JOIN nhanvien nv ON cvnn.nhan_vien_id = nv.id
          JOIN du_an d ON cv.du_an_id = d.id
          WHERE d.ten_du_an LIKE '%Web HRM%'         

- User: "Hôm nay ai đang nghỉ phép?" 
  -> SQL: SELECT nv.ho_ten, dnp.ly_do FROM don_nghi_phep dnp JOIN nhanvien nv ON dnp.nhan_vien_id = nv.id WHERE CURRENT_DATE BETWEEN dnp.ngay_bat_dau AND dnp.ngay_ket_thuc AND dnp.trang_thai = 'da_duyet'
- User: "Nguyễn Tấn Dũng còn bao nhiêu phép?"
  -> SQL: SELECT nv.ho_ten, np.ngay_phep_con_lai FROM ngay_phep_nam np JOIN nhanvien nv ON np.nhan_vien_id = nv.id WHERE nv.ho_ten LIKE '%Nguyễn Tấn Dũng%' AND np.nam = YEAR(CURRENT_DATE)
- User: "Nhân viên Nguyễn Tấn Dũng thuộc phòng ban nào?"
  -> SQL: SELECT nv.ho_ten, pb.ten_phong AS phong_ban FROM nhanvien nv JOIN phong_ban pb ON nv.phong_ban_id = pb.id WHERE nv.ho_ten LIKE '%Nguyễn Tấn Dũng%' AND nv.trang_thai_lam_viec != 'Nghỉ việc'
- User: "Thông tin phép năm của nhân viên có id = 5"
  -> SQL: SELECT nv.ho_ten, np.tong_ngay_phep, np.ngay_phep_da_dung, np.ngay_phep_con_lai FROM ngay_phep_nam np JOIN nhanvien nv ON np.nhan_vien_id = nv.id WHERE nv.id = 5 AND np.nam = YEAR(CURRENT_DATE)
- User: "Tổng phép của toàn công ty năm 2026?"
  -> SQL: SELECT COUNT(DISTINCT np.nhan_vien_id) as so_nhan_vien, SUM(np.tong_ngay_phep) as tong_phep FROM ngay_phep_nam np WHERE np.nam = 2026

- User: "Công việc 205 có bao nhiêu bước?"
  -> SQL: SELECT COUNT(*) as so_buoc FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND trang_thai != 'Hủy'
- User: "Liệt kê các bước của công việc 205"
  -> SQL: SELECT ten_buoc, trang_thai, ngay_bat_dau, ngay_ket_thuc FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND trang_thai != 'Hủy' ORDER BY ngay_bat_dau
- User: "Các bước nào của công việc 205 đã hoàn thành?"
  -> SQL: SELECT ten_buoc, ngay_ket_thuc FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND trang_thai = 'Đã hoàn thành'
- User: "Bước nào của công việc 205 trễ hạn?"
  -> SQL: SELECT ten_buoc, ngay_ket_thuc FROM cong_viec_quy_trinh WHERE cong_viec_id = 205 AND ngay_ket_thuc < CURRENT_DATE AND trang_thai != 'Hủy' AND trang_thai != 'Đã hoàn thành'
- User: "Tìm bước 'Phân tích yêu cầu' trong công việc"
  -> SQL: SELECT ten_buoc, trang_thai, mo_ta FROM cong_viec_quy_trinh WHERE ten_buoc LIKE '%Phân tích%'
- User: "Công việc 'Oracle Cloud' có bao nhiêu bước?"
  -> SQL: SELECT COUNT(*) as so_buoc FROM cong_viec_quy_trinh WHERE cong_viec_id IN (SELECT id FROM cong_viec WHERE ten_cong_viec LIKE '%Oracle Cloud%' AND trang_thai != 'Hủy') AND trang_thai != 'Hủy'
- User: "Liệt kê bước của công việc 'Gurucul AI Agent'"
  -> SQL: SELECT ten_buoc, trang_thai, ngay_bat_dau, ngay_ket_thuc FROM cong_viec_quy_trinh WHERE cong_viec_id IN (SELECT id FROM cong_viec WHERE ten_cong_viec LIKE '%Gurucul%') AND trang_thai != 'Hủy'

- User: "Danh sách nhân viên còn dưới 5 ngày phép?"
  -> SQL: SELECT nv.ho_ten, np.ngay_phep_con_lai FROM ngay_phep_nam np JOIN nhanvien nv ON np.nhan_vien_id = nv.id WHERE np.nam = YEAR(CURRENT_DATE) AND np.ngay_phep_con_lai < 5
- User: "Nhân viên nào đã dùng hết phép?"
  -> SQL: SELECT nv.ho_ten, np.ngay_phep_con_lai FROM ngay_phep_nam np JOIN nhanvien nv ON np.nhan_vien_id = nv.id WHERE np.nam = YEAR(CURRENT_DATE) AND np.ngay_phep_con_lai <= 0
- User: "Giám đốc công ty là ai?" -> SQL: SELECT ho_ten, chuc_vu, email, so_dien_thoai FROM nhanvien WHERE chuc_vu LIKE '%Giám đốc%' OR chuc_vu LIKE '%CEO%' OR chuc_vu LIKE '%General Manager%'
- User: "Chấm công hôm 04/02/2026 của Trần Đình Nam" -> SQL: SELECT c.ngay, c.check_in, c.check_out FROM cham_cong c JOIN nhanvien n ON c.nhan_vien_id = n.id WHERE c.ngay = '2026-02-04' AND n.ho_ten LIKE '%Trần Đình Nam%'
- User: "Ai đi muộn hôm nay?" -> SQL: SELECT nv.ho_ten, c.check_in FROM cham_cong c JOIN nhanvien nv ON c.nhan_vien_id = nv.id WHERE c.ngay = CURRENT_DATE AND c.check_in >= '08:06:00'
- User: "Danh sách chấm công hôm 2026-02-04" -> SQL: SELECT nv.ho_ten, c.check_in, c.check_out FROM cham_cong c JOIN nhanvien nv ON c.nhan_vien_id = nv.id WHERE c.ngay = '2026-02-04'
- User: "Ai vắng mặt hôm nay?" -> SQL: SELECT ho_ten FROM nhanvien WHERE id NOT IN (SELECT nhan_vien_id FROM cham_cong WHERE ngay = CURRENT_DATE) AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Bao nhiêu nhân viên đi muộn tháng 1/2026?" -> SQL: SELECT COUNT(DISTINCT c.nhan_vien_id) AS so_nhan_vien_di_muon FROM cham_cong c WHERE c.check_in >= '08:06:00' AND MONTH(c.ngay) = 1 AND YEAR(c.ngay) = 2026
- User: "Bao nhiêu lần đi muộn tháng 1/2026?" -> SQL: SELECT COUNT(*) AS tong_so_lan_di_muon FROM cham_cong c WHERE c.check_in >= '08:06:00' AND MONTH(c.ngay) = 1 AND YEAR(c.ngay) = 2026
- User: "Tổng nhân viên vắng mặt tháng 1/2026?" -> SQL: SELECT COUNT(DISTINCT nhanvien.id) FROM nhanvien WHERE nhanvien.id NOT IN (SELECT DISTINCT nhan_vien_id FROM cham_cong WHERE MONTH(ngay) = 1 AND YEAR(ngay) = 2026) AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Lấy thông tin nhân viên có id = 1" -> SQL: SELECT ho_ten, email, so_dien_thoai, chuc_vu FROM nhanvien WHERE id = 1 AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Ngày sinh của nhân viên có id = 7" -> SQL: SELECT ngay_sinh FROM nhanvien WHERE id = 7 AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Giới tính của nhân viên Đình Nam" -> SQL: SELECT gioi_tinh FROM nhanvien WHERE ho_ten LIKE '%Đình Nam%' AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Email của nhân viên có id = 5" -> SQL: SELECT email FROM nhanvien WHERE id = 5 AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Số điện thoại của nhân viên tên Trần Minh" -> SQL: SELECT so_dien_thoai FROM nhanvien WHERE ho_ten LIKE '%Trần Minh%' AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Nhân viên nào có thâm niên lâu nhất?" -> SQL: SELECT ho_ten, ngay_vao_lam FROM nhanvien WHERE trang_thai_lam_viec != 'Nghỉ việc' ORDER BY ngay_vao_lam ASC LIMIT 1
- User: "Nhân viên nào vào làm gần đây nhất?" -> SQL: SELECT ho_ten, ngay_vao_lam FROM nhanvien WHERE trang_thai_lam_viec != 'Nghỉ việc' ORDER BY ngay_vao_lam DESC LIMIT 1
- User: "Tìm nhân viên tên Nguyễn Văn A" -> SQL: SELECT ho_ten, email, so_dien_thoai, chuc_vu FROM nhanvien WHERE ho_ten LIKE '%Nguyễn Văn A%' AND trang_thai_lam_viec != 'Nghỉ việc'
- User: "Danh sách toàn bộ nhân viên" -> SQL: SELECT ho_ten FROM nhanvien WHERE trang_thai_lam_viec != 'Nghỉ việc'
- User: "Có bao nhiêu nhân viên?" -> SQL: SELECT COUNT(*) as tong_so_nhan_vien FROM nhanvien WHERE trang_thai_lam_viec != 'Nghỉ việc'
- User: "Danh sách nhân viên đang làm việc" -> SQL: SELECT ho_ten FROM nhanvien WHERE trang_thai_lam_viec = 'Đang làm'
- User: "Danh sách nhân viên đã nghỉ việc" -> SQL: SELECT ho_ten FROM nhanvien WHERE trang_thai_lam_viec = 'Nghỉ việc'

- User: "Danh sách công việc đang thực hiện" -> SQL: SELECT ten_cong_viec, han_hoan_thanh FROM cong_viec WHERE trang_thai = 'Đang thực hiện' AND trang_thai != 'Hủy'
- User: "Công việc nào trễ hạn?" -> SQL: SELECT ten_cong_viec, han_hoan_thanh FROM cong_viec WHERE han_hoan_thanh < CURRENT_DATE AND trang_thai NOT IN ('Đã hoàn thành', 'Hủy')
- User: "Tôi cần tìm info công việc 'Oracle Cloud'" -> SQL: SELECT ten_cong_viec, trang_thai, han_hoan_thanh FROM cong_viec WHERE ten_cong_viec LIKE '%Oracle Cloud%'
- User: "Bao nhiêu công việc được giao cho nhân viên 'Trần Minh'?" -> SQL: SELECT COUNT(DISTINCT cv.id) FROM cong_viec cv JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id JOIN nhanvien nv ON cvnn.nhan_vien_id = nv.id WHERE nv.ho_ten LIKE '%Trần Minh%' AND cv.trang_thai != 'Hủy'

- User: "Có bao nhiêu cấu hình hệ thống?" -> SQL: SELECT COUNT(*) as tong_so_cau_hinh FROM cau_hinh_he_thong
- User: "Danh sách cấu hình hệ thống" -> SQL: SELECT ten_cau_hinh, gia_tri FROM cau_hinh_he_thong
- User: "Cấu hình nào liên quan tới email?" -> SQL: SELECT ten_cau_hinh, gia_tri, mo_ta FROM cau_hinh_he_thong WHERE ten_cau_hinh LIKE '%email%'
- User: "Cấu hình working_hours được tạo bao giờ?" -> SQL: SELECT ngay_tao FROM cau_hinh_he_thong WHERE ten_cau_hinh LIKE '%working_hours%'

- User: "Có bao nhiêu sự kiện trong hệ thống?" -> SQL: SELECT COUNT(*) as tong_su_kien FROM lich_trinh
- User: "Liệt kê tất cả sự kiện" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh ORDER BY ngay_bat_dau DESC
- User: "Sự kiện nào hôm nay?" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE ngay_bat_dau = CURRENT_DATE ORDER BY ngay_bat_dau
- User: "Sự kiện sắp tới trong tuần tới?" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE ngay_bat_dau >= CURRENT_DATE AND ngay_bat_dau < DATE_ADD(CURRENT_DATE, INTERVAL 7 DAY) ORDER BY ngay_bat_dau
- User: "Sự kiện nào đã qua?" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE ngay_ket_thuc < CURRENT_DATE ORDER BY ngay_ket_thuc DESC
- User: "Tìm sự kiện 'Hội thảo'" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE tieu_de LIKE '%Hội thảo%' ORDER BY ngay_bat_dau
- User: "Sự kiện 'LBS startup showcase' kết thúc khi nào?" -> SQL: SELECT tieu_de, ngay_ket_thuc FROM lich_trinh WHERE tieu_de LIKE '%LBS startup showcase%'
- User: "Tất cả sự kiện được tạo trong tháng 2/2026?" -> SQL: SELECT tieu_de, ngay_bat_dau, ngay_tao FROM lich_trinh WHERE MONTH(ngay_tao) = 2 AND YEAR(ngay_tao) = 2026 ORDER BY ngay_bat_dau
- User: "Sự kiện nào có ngày bắt đầu là 2026-02-15?" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE ngay_bat_dau = '2026-02-15'
- User: "Liệt kê sự kiện kéo dài nhiều ngày (từ 2 ngày trở lên)" -> SQL: SELECT tieu_de, ngay_bat_dau, ngay_ket_thuc, DATEDIFF(ngay_ket_thuc, ngay_bat_dau) as so_ngay FROM lich_trinh WHERE DATEDIFF(ngay_ket_thuc, ngay_bat_dau) > 1 ORDER BY ngay_bat_dau
- User: "Danh sách sự kiện sắp tới sắp xếp theo ngày?" -> SQL: SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc FROM lich_trinh WHERE ngay_bat_dau >= CURRENT_DATE ORDER BY ngay_bat_dau ASC

- User: "Có bao nhiêu quyền hạn?" -> SQL: SELECT COUNT(*) as tong_quyen FROM quyen
- User: "Liệt kê tất cả quyền hạn" -> SQL: SELECT id, ma_quyen, ten_quyen, nhom_quyen FROM quyen ORDER BY nhom_quyen, ma_quyen
- User: "Quyền hạn liên quan tới phòng ban?" -> SQL: SELECT ma_quyen, ten_quyen FROM quyen WHERE nhom_quyen LIKE '%phongban%'
- User: "Tìm quyền 'Thêm phòng ban'" -> SQL: SELECT id, ma_quyen, ten_quyen FROM quyen WHERE ten_quyen LIKE '%Thêm phòng ban%'
- User: "Danh sách quyền hạn của nhóm nhanvien" -> SQL: SELECT ma_quyen, ten_quyen FROM quyen WHERE nhom_quyen = 'nhanvien' ORDER BY ten_quyen
- User: "Có bao nhiêu nhóm quyền khác nhau?" -> SQL: SELECT COUNT(DISTINCT nhom_quyen) as so_nhom_quyen FROM quyen
- User: "Liệt kê tất cả nhóm quyền" -> SQL: SELECT DISTINCT nhom_quyen FROM quyen ORDER BY nhom_quyen
- User: "Bao nhiêu quyền hạn trong nhóm admin?" -> SQL: SELECT COUNT(*) as so_quyen_admin FROM quyen WHERE nhom_quyen = 'admin'
- User: "Quyền nào liên quan tới 'quản lý'?" -> SQL: SELECT id, ma_quyen, ten_quyen, nhom_quyen FROM quyen WHERE ten_quyen LIKE '%quản lý%' OR ten_quyen LIKE '%Quản lý%'

- User: "Có bao nhiêu tài liệu?" -> SQL: SELECT COUNT(*) as tong_tai_lieu FROM tai_lieu WHERE trang_thai != 'Đã xóa'
- User: "Liệt kê tất cả tài liệu" -> SQL: SELECT ten_tai_lieu, loai_tai_lieu, file_size, luot_xem, ngay_tao FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY ngay_tao DESC
- User: "Tài liệu nào được xem nhiều nhất?" -> SQL: SELECT ten_tai_lieu, loai_tai_lieu, luot_xem, file_name FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY luot_xem DESC LIMIT 1
- User: "Tài liệu nào được tải nhiều nhất?" -> SQL: SELECT ten_tai_lieu, loai_tai_lieu, luot_tai, file_size, file_name FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY luot_tai DESC LIMIT 1
- User: "Tìm tài liệu 'File cài đặt CSA'" -> SQL: SELECT ten_tai_lieu, mo_ta, file_name, file_size, loai_tai_lieu, ngay_tao FROM tai_lieu WHERE ten_tai_lieu LIKE '%File cài đặt CSA%'
- User: "Danh sách tài liệu loại 'Báo cáo'" -> SQL: SELECT ten_tai_lieu, mo_ta, file_size, luot_xem, luot_tai FROM tai_lieu WHERE loai_tai_lieu = 'Báo cáo' AND trang_thai != 'Đã xóa' ORDER BY ngay_tao DESC
- User: "Tài liệu nào được tạo trong tháng 1/2026?" -> SQL: SELECT ten_tai_lieu, loai_tai_lieu, file_name, ngay_tao FROM tai_lieu WHERE trang_thai != 'Đã xóa' AND MONTH(ngay_tao) = 1 AND YEAR(ngay_tao) = 2026 ORDER BY ngay_tao DESC
- User: "Bao nhiêu tài liệu cho phép xem 'Tất cả'?" -> SQL: SELECT COUNT(*) as so_tai_lieu FROM tai_lieu WHERE trang_thai != 'Đã xóa' AND doi_tuong_xem = 'Tất cả'
- User: "Danh sách tài liệu do nhân viên ID 18 tạo" -> SQL: SELECT ten_tai_lieu, loai_tai_lieu, file_size, luot_xem, ngay_tao FROM tai_lieu WHERE nguoi_tao_id = 18 AND trang_thai != 'Đã xóa' ORDER BY ngay_tao DESC
- User: "Tài liệu nào có kích thước lớn nhất?" -> SQL: SELECT ten_tai_lieu, file_name, file_size, loai_tai_lieu, ngay_tao FROM tai_lieu WHERE trang_thai != 'Đã xóa' ORDER BY file_size DESC LIMIT 1

⚠️ CÁC GIÁ TRỊ CHÍNH XÁC CỦA TRẠNG THÁI LÀM VIỆC:
- 'Đang làm' = Nhân viên đang làm việc
- 'Nghỉ việc' = Nhân viên đã nghỉ việc
- TUYỆT ĐỐI KHÔNG dùng 'Đang làm việc' (sai!)
- Khi lọc nhân viên đang hoạt động, dùng: WHERE trang_thai_lam_viec = 'Đang làm'
- Khi loại bỏ nhân viên nghỉ việc, dùng: WHERE trang_thai_lam_viec != 'Nghỉ việc'

⚠️ LƯU Ý QUAN TRỌNG VỀ SELECT * & CÁCH XỬ LÝ DANH SÁCH:
- Tuyệt đối CẤM dùng SELECT * (chỉ select những cột cần thiết)
- Câu hỏi "danh sách nhân viên" -> Chỉ lấy ho_ten, KHÔNG lấy tất cả cột
- Câu hỏi yêu cầu "tất cả thông tin" mới dùng SELECT * hoặc liệt kê tường minh các cột
- Áp dụng tương tự cho các bảng khác (dự án, công việc, ...)
- MẶC ĐỊNH: Khi query nhân viên (danh sách, thông tin cá nhân, tìm kiếm) -> Loại bỏ nhân viên đã nghỉ việc bằng AND trang_thai_lam_viec != 'Nghỉ việc'
- NGOẠI LỆ: Chỉ bỏ điều kiện này nếu người dùng tường minh hỏi về nhân viên đã nghỉ việc

SCHEMA:
{schema}

CÂU HỎI:
{question}

SQL OUTPUT (Only SQL):
""")

# --- PROMPT 2: ĐỌC BÁO CÁO (Humanize Answer) ---
ANSWER_PROMPT = ChatPromptTemplate.from_template("""
Bạn là trợ lý HRM thông minh.
Nhiệm vụ: Đọc dữ liệu JSON và trả lời câu hỏi của người dùng.

THÔNG TIN:
- Câu hỏi: "{question}"
- Dữ liệu nhận được: {data}

⚠️ **NHẬP CUỐI CÙNG - PHẢI ĐỌC TRƯỚC HẾT:**
- Nếu dữ liệu là `[]` (empty list) hoặc `null` hoặc "None" → Hãy trả lời một câu đầy đủ, lịch sự (ví dụ: "Hiện tại không có [cái gì đó]."), không suy luận thêm lý do ẩn phía sau
- Nếu dữ liệu có ít nhất 1 dòng → PHẢI liệt kê tất cả, không được tóm tắt

YÊU CẦU TRẢ LỜI:

1. Nếu dữ liệu KHÔNG rỗng (có ít nhất 1 dòng):
  - Viết câu trả lời theo văn phong tự nhiên, đầy đủ, không quá cộc lốc.
  - Có thể mở đầu bằng 1 câu tóm tắt ngắn (ví dụ: "Dưới đây là danh sách kết quả:"), sau đó liệt kê chi tiết.
  - **BẮT BUỘC LIỆT KÊ CHI TIẾT tất cả dòng dữ liệu**, KHÔNG ĐƯỢC tóm tắt
  - Hiển thị tất cả cột từ SQL query
  - Mỗi bản ghi phải có tên/ID + các thông tin khác đầy đủ
  - Dùng dấu "-" hoặc bullet point cho từng mục
   - **⚠️ CRITICAL: LIỆT KÊ TỪNG DÒNG MỘT - KHÔNG BỎ QUA DÒNG NÀO:**
     + Nếu SQL trả về 5 dòng → PHẢI in ra 5 dòng (không được chỉ in 2-3 dòng)
     + Nếu SQL trả về 20 dòng → PHẢI in ra 20 dòng (không được chỉ in 10 dòng)
     + KHÔNG được nói "và những cái khác..." hoặc "v.v." ở cuối
     + PHẢI LIỆT KÊ HẾT tất cả, đếm từng dòng được
   - **KHÔNG ĐƯỢC:** Nói "Có X cái..." rồi dừng (VD: "Có 2 dự án trễ hạn")
   - **PHẢI:** Liệt kê chi tiết tên + ngày + trạng thái + v.v. của TỪNG dự án
   - **KHÔNG ĐƯỢC suy luận thêm** dựa trên logic (VD: "không có dự án vì ngày kết thúc trong tương lai")
   - **CHỈ báo cáo lại dữ liệu SQL mà thôi**

2. Nếu dữ liệu rỗng (Empty List hoặc Null hoặc "[]"):
   
   A. CÂUỎI TÌM KIẾM CỤ THỂ (By ID, By Name):
   - Ví dụ: "Lấy thông tin nhân viên có id = 1", "Tìm nhân viên tên Nguyễn Văn A"
   - Nếu không có dữ liệu → Báo THỰC TẾ: "Không tìm thấy nhân viên với điều kiện này"
   - KHÔNG được suy đoán hay nói "chưa được thêm vào hệ thống"
   
   B. CÂUHỎI KIỂM TRA TRẠNG THÁI HÔM NAY (Chỉ áp dụng cho chấm công hôm nay):
   - Ví dụ: "Hôm nay ai đi muộn?", "Hôm nay ai vắng mặt?", "Hôm nay ai nghỉ phép?"
   - Điều kiện: Phải có từ khóa "hôm nay" hoặc chúng ta có bối cảnh ngày cụ thể
   - Nếu không có dữ liệu → Được phép suy luận tích cực:
     + "Hôm nay không có nhân viên nào đi muộn."
     + "Hôm nay toàn bộ nhân viên đều đi làm đầy đủ."
   
   C. CÂUHỎI VỀ DANH SÁCH & THỐNG KÊ CHUNG (Không hạn chế theo thời gian):
   - Ví dụ: "Nhân viên nào tạm dừng công việc?", "Danh sách công việc trễ hạn", "Dự án trễ hạn", "Ai chưa nộp báo cáo?"
   - Nếu không có dữ liệu → Báo THỰC TẾ: "Không có nhân viên nào tạm dừng công việc"
   - ⚠️ **CẢNH BÁO VỀ CÂU HỎI "DỰ ÁN/CÔNG VIỆC TRỄ HẠN":**
     + **TUYỆT ĐỐI KHÔNG suy luận** "không có dự án nào trễ hạn vì tất cả ngày kết thúc đều trong tương lai"
     + **Nếu SQL trả về dữ liệu** → Báo danh sách dự án đó (đây chính là dự án trễ hạn)
     + **Nếu SQL trả về rỗng** → Báo "Không có dự án nào trễ hạn" (không giải thích lý do)
     + Đừng suy đoán, chỉ báo cáo dữ liệu thực
   - KHÔNG được suy đoán, dữ liệu là sự thật

3. Với dữ liệu thống kê (COUNT, SUM, AVG):
   - Nếu dữ liệu là một con số, đó chính là câu trả lời
   - Trả lời trực tiếp, không nói thiếu thông tin

4. Khi SQL đã có điều kiện lọc:
   - Mặc định TẤT CẢ bản ghi trả về đều thỏa mãn điều kiện
   - Không cần suy đoán thêm từ phía AI

5. TRUNG THỰC VỚI DỮ LIỆU (DATA FIDELITY – BẮT BUỘC):
   - Không được tự ý loại bỏ bất kỳ bản ghi nào
   - Không được bỏ qua các giá trị 0 (0% tiến độ là thông tin hợp lệ)
   - SQL trả về gì → câu trả lời phải phản ánh đúng như vậy
   - **⚠️ CẢNH BÁO CỰC CÓT LÕI: KHÔNG ĐƯỢC SỬ DỤNG LOGIC ĐỂ SỬ LÝ DỮ LIỆU RỖNG**
     + KHÔNG nói "Không có dự án trễ hạn vì tất cả ngày kết thúc đều trong tương lai"
     + KHÔNG giải thích lý do tại sao dữ liệu rỗng
     + CHỈ báo: "Không có dự án nào trễ hạn" (nếu dữ liệu thực sự rỗng)
     + NẾU dữ liệu SQL trả về có dòng → PHẢI liệt kê, không được nói "không có"

6. QUY TẮC ĐỊNH DẠNG (BẮT BUỘC):
  - TUYỆT ĐỐI KHÔNG dùng Markdown in đậm (**).
  - KHÔNG dùng **text** trong mọi trường hợp.
  - Chỉ trả lời bằng văn bản thường.
  - Nếu cần liệt kê → dùng dấu "-" ở đầu dòng.

7. QUY TẮC ĐỊNH DẠNG DỮ LIỆU CHẤM CÔNG (ATTENDANCE DATA FORMAT):
   - Khi dữ liệu là bảng chấm công (cham_cong), PHẢI định dạng theo nhóm nhân viên.
   - **KHÔNG** mỗi dòng chấm công đều gạch đầu dòng riêng biệt.
   - **CÓ** gom nhóm theo Nhân viên ID/Tên rồi mới gạch đầu dòng cho từng bản ghi chấm công.
   - Ví dụ ĐÚNG:
     Nhân viên Trần Minh (ID: 5):
     - Ngày 2026-02-04: Check in 08:15, Check out 17:30
     - Ngày 2026-02-03: Check in 08:00, Check out 17:00
     Nhân viên Nguyễn Văn A (ID: 3):
     - Ngày 2026-02-04: Check in 08:06 (Đi muộn), Check out 17:00
     - Ngày 2026-02-03: Check in 08:00, Check out 16:45
   
   - Ví dụ SAI (không làm):
     - ID 5, Ngày 2026-02-04, Check in 08:15, Check out 17:30
     - ID 5, Ngày 2026-02-03, Check in 08:00, Check out 17:00
     - ID 3, Ngày 2026-02-04, Check in 08:06, Check out 17:00

8. QUY TẮC ĐỊNH DẠNG DỮ LIỆU PHÉP NĂM (ANNUAL LEAVE DATA FORMAT - RẤT QUAN TRỌNG):
   - Khi dữ liệu là bảng phép năm (ngay_phep_nam), PHẢI group theo nhân viên.
   - Nếu có NHIỀU nhân viên (kết quả nhiều dòng) → Liệt kê từng nhân viên rõ ràng
   - Nếu có MỘT nhân viên (kết quả 1 dòng) → Trực tiếp hiển thị thông tin
   - Ý nghĩa các trường phổ biến:
     + `tong_ngay_phep`: tổng ngày phép hiện tại của nhân viên trong năm (do hệ thống tính, có thể đã bao gồm phép năm trước hoặc cộng đầu năm).
     + `ngay_phep_da_dung`: số ngày phép đã sử dụng.
     + `ngay_phep_con_lai`: số ngày phép còn lại (nếu có cột này trong dữ liệu).
     + Có thể tồn tại thêm các trường khác như `ngay_phep_nam_truoc`, `da_cong_phep_dau_nam`, ... → chỉ cần đọc đúng giá trị.
   - KHÔNG tự đặt công thức cứng giữa các cột. Hãy coi các con số trong dữ liệu là sự thật hệ thống và trình bày lại cho người dùng (tổng, đã dùng, còn lại, phép năm trước, v.v.) theo đúng ý nghĩa trường.

GIỌNG ĐIỆU:
Tự nhiên, thân thiện, chuyên nghiệp, giống trợ lý nội bộ doanh nghiệp.

TRẢ LỜI:
""")


# --- PROMPT 3: TRẢ LỜI TỪ TÀI LIỆU TĨNH (STATIC DOCX RAG) ---
STATIC_ANSWER_PROMPT = ChatPromptTemplate.from_template("""
Bạn là trợ lý tư vấn giải pháp nội bộ của ICS.
Nhiệm vụ: Trả lời câu hỏi dựa trên NỘI DUNG TÀI LIỆU dưới đây.

VĂN BẢN THAM KHẢO (trích từ tài liệu nội bộ):
{context}

CÂU HỎI:
{question}

YÊU CẦU:
- Chỉ sử dụng thông tin có trong văn bản tham khảo ở trên.
- Không bịa thêm, không suy luận ngoài nội dung tài liệu.
- Trả lời chi tiết, rõ ràng, bao quát đầy đủ các ý liên quan đến câu hỏi.
- Có thể dùng nhiều câu hoặc gạch đầu dòng nếu cần, không tự ý rút gọn hay lược bỏ các ý quan trọng trong tài liệu.
- Không lặp lại nguyên văn toàn bộ đoạn văn rất dài; chỉ tóm tắt phần liên quan nhưng vẫn giữ đủ nội dung cần thiết để người đọc hiểu rõ.
- Nếu trong tài liệu KHÔNG có thông tin để trả lời → nói rõ: "Tôi chưa có dữ lieu cho nội dung bạn vừa đề cập. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý nhân sự (HRM) bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?".
 - Tuyệt đối KHÔNG dùng Markdown in đậm (**), không bao quanh tiêu đề hoặc cụm từ bằng cặp ** ở hai bên.
 - Chỉ trả lời bằng văn bản thường; nếu cần nhấn mạnh có thể dùng câu chữ, không dùng định dạng Markdown.

TRẢ LỜI:
""")


# ==========================================================
# 4.5. DOWNLOAD FILE ENDPOINT
# ==========================================================
from fastapi.responses import FileResponse
import os

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Serve exported files (docx/pdf) for download"""
    # Security: Validate filename to prevent directory traversal
    if "../" in filename or "..\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    
    filepath = os.path.join(EXPORT_DIR, filename)
    
    # Check if file exists
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Return file for download
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ==========================================================
# 4. HELPER FUNCTIONS (Xử lý & Gọi API)
# ==========================================================
def validate_sql(sql: str) -> str:
    """Làm sạch và kiểm tra an toàn SQL"""
    # Xóa markdown nếu có
    sql_clean = sql.replace("```sql", "").replace("```", "").strip()
    
    # Chặn các lệnh nguy hiểm (Chỉ cho phép SELECT)
    forbidden = ["insert", "update", "delete", "drop", "alter", "truncate", "grant"]
    if any(cmd in sql_clean.lower() for cmd in forbidden):
        print(f"⚠️ Blocked dangerous SQL: {sql_clean}")
        return ""
    
    return sql_clean

def execute_sql_api(sql: str) -> Any:
    """Gọi API HRM để lấy dữ liệu"""
    if not sql: return None

    # Log query ra terminal để debug
    print(f"\n[DEBUG SQL]: {sql}")

    try:
        payload = {"command": sql}
        res = requests.post(HRM_API_URL, json=payload, timeout=30)
        
        if res.status_code == 200:
            try:
                # Ưu tiên trả về JSON object
                return res.json()
            except:
                return res.text
        else:
            print(f"❌ API Error {res.status_code}: {res.text}")
            return f"Lỗi từ hệ thống dữ liệu: {res.text}"
    except Exception as e:
        print(f"❌ Connection Error: {e}")
        return "Lỗi kết nối đến máy chủ dữ liệu."


def validate_leave_balance_data(data_result: Any) -> tuple[bool, str]:
    """
    Validate dữ liệu phép năm (ngay_phep_nam).
  Công thức tham chiếu: tong_ngay_phep ≈ ngay_phep_da_dung + ngay_phep_con_lai
  LƯU Ý: Hiện tại chỉ cảnh báo (log) nếu lệch, KHÔNG chặn trả lời.
  Returns: (is_valid, error_message)
    """
    if not data_result or isinstance(data_result, str):
        return (True, "")  # Không validate nếu không có dữ liệu
    
    try:
        # Nếu là list
        if isinstance(data_result, list):
            for record in data_result:
                if isinstance(record, dict):
                    if 'tong_ngay_phep' in record and 'ngay_phep_da_dung' in record and 'ngay_phep_con_lai' in record:
                        total = float(record.get('tong_ngay_phep', 0))
                        used = float(record.get('ngay_phep_da_dung', 0))
                        remaining = float(record.get('ngay_phep_con_lai', 0))
                        
                        # Kiểm tra công thức (soft check - chỉ cảnh báo, không chặn)
                        expected_total = used + remaining
                        if abs(total - expected_total) > 0.01:  # Cho phép sai số 0.01
                          print(f"[WARN] Dữ liệu phép năm có thể không hợp lệ (Tổng: {total} ≠ Đã dùng {used} + Còn lại {remaining} = {expected_total}).")
        
        # Nếu là single dict
        elif isinstance(data_result, dict):
            if 'tong_ngay_phep' in data_result and 'ngay_phep_da_dung' in data_result and 'ngay_phep_con_lai' in data_result:
                total = float(data_result.get('tong_ngay_phep', 0))
                used = float(data_result.get('ngay_phep_da_dung', 0))
                remaining = float(data_result.get('ngay_phep_con_lai', 0))
                
                expected_total = used + remaining
                if abs(total - expected_total) > 0.01:
                  print(f"[WARN] Dữ liệu phép năm có thể không hợp lệ (Tổng: {total} ≠ Đã dùng {used} + Còn lại {remaining} = {expected_total}).")
        
        return (True, "")
    except Exception as e:
        print(f"Error validating leave balance: {e}")
        return (True, "")  # Nếu lỗi parse, bỏ qua validation

# ==========================================================
# KIỂM TRA CÂU HỎI NGOÀI LỀ
# ==========================================================
def is_out_of_scope_question(question: str) -> bool:
    """Phát hiện câu hỏi hoàn toàn ngoài phạm vi HRM"""
    q = question.lower().strip()
    
    # Từ khóa NGOÀI LỀ - hoàn toàn không liên quan HRM
    out_of_scope = [
        'thời tiết', 'mưa', 'nắng', 'gió', 'bão', 'nhiệt độ',
        'bóng đá', 'bóng chuyền', 'bóng rổ', 'cầu lông', 'tennis', 'vđv',
        'chính trị', 'bầu cử', 'tổng thống', 'thủ tướng', 'chiến tranh',
        'phim', 'nhạc sĩ', 'ca sĩ', 'diễn viên', 'gameshow',
        'cổ phiếu', 'chứng khoán', 'tiền',
        'bệnh', 'dị ứng', 'khám bệnh', 'thuốc',
        'tôn giáo', 'phật giáo', 'hồi giáo',
        'bạn là ai', 'bạn tên gì', 'bạn ở đâu', 'bạn bao nhiêu tuổi',
        'code', 'lập trình', 'python', 'javascript', 'docker',
    ]
    
    return any(word in q for word in out_of_scope)

# ==========================================================
# 5. MAIN ENDPOINT (Luồng xử lý chính)
# ==========================================================
@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(req: ChatRequest):
    try:
        question_lower = req.question.lower()
        
        # KIỂM TRA TRƯỚC: Câu hỏi ngoài lề
        if is_out_of_scope_question(req.question):
            return ChatResponse(
                sql=None,
                data=None,
                answer="Tôi là chatbot chuyên về Quản lý Nhân sự (HRM). Câu hỏi của bạn nằm ngoài phạm vi hỗ trợ. Bạn có câu hỏi nào về nhân sự, dự án, hoặc công việc không?",
                download_url=None
            )
        
        # KIỂM TRA TRƯỚC: Câu hỏi về lương (dữ liệu nhạy cảm)
        salary_keywords = ['lương', 'mức lương', 'tiền lương', 'hệ số lương', 'lương cơ bản']
        
        if any(keyword in question_lower for keyword in salary_keywords):
            return ChatResponse(
                sql=None,
                data=None,
                answer="Dữ liệu lương là thông tin cá nhân và nhạy cảm, không được công khai. Tôi không thể truy vấn thông tin này. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý nhân sự (HRM) bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?",
                download_url=None
            )
        
        # KIỂM TRA TRƯỚC: Câu hỏi về KPI (dữ liệu không công khai)
        kpi_keywords = ['kpi', 'performance', 'target']
        
        if any(keyword in question_lower for keyword in kpi_keywords):
            return ChatResponse(
                sql=None,
                data=None,
                answer="Dữ liệu KPI là thông tin không công khai. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý nhân sự (HRM) bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?",
                download_url=None
            )
        
        # KIỂM TRA TRƯỚC: Câu hỏi về lịch sử nhân sự (dữ liệu không công khai)
        history_keywords = ['lịch sử', 'quá khứ', 'công việc trước', 'kinh nghiệm làm việc', 'hồ sơ', 'tiểu sử', 'trước đây',]
        
        if any(keyword in question_lower for keyword in history_keywords):
            return ChatResponse(
                sql=None,
                data=None,
                answer="Dữ liệu lịch sử nhân sự là thông tin không được công khai. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý nhân sự (HRM) bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?",
                download_url=None
            )

        # ROUTER: quyết định ưu tiên HRM_SQL hay STATIC_DOC
        route = "HRM_SQL"
        try:
          route_raw = router_chain.invoke({"question": req.question})
          route = (route_raw or "").strip().upper()
        except Exception as e:
          print(f"[Router] error: {e}")
          route = "HRM_SQL"

        if route not in ("HRM_SQL", "STATIC_DOC"):
          route = "HRM_SQL"

        # Nếu router chọn STATIC_DOC → bỏ qua SQL, trả lời từ tài liệu tĩnh
        if route == "STATIC_DOC":
          static_context = build_static_context(req.question, k=3)
          if static_context:
            static_chain = STATIC_ANSWER_PROMPT | llm_answer | StrOutputParser()
            static_answer = static_chain.invoke({
              "question": req.question,
              "context": static_context,
            })
            # Làm sạch Markdown in đậm nếu vẫn còn
            if isinstance(static_answer, str):
                static_answer = static_answer.replace("**", "")
            return ChatResponse(sql=None, data=None, answer=static_answer, download_url=None)
          # Không tìm được gì trong tài liệu tĩnh
          return ChatResponse(
            sql=None,
            data=None,
            answer="Tôi chưa có dữ lieu cho nội dung bạn vừa đề cập. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý nhân sự (HRM) bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?",
            download_url=None,
          )

        # KIỂM TRA TRƯỚC: Câu hỏi về chấm công/đi muộn ngày tương lai
        from datetime import datetime, timedelta
        import re
        
        # Từ khóa liên quan chấm công
        attendance_keywords = ['chấm công', 'chấm công', 'đi muộn', 'check in', 'check out', 'vắng mặt', 'không đi làm']
        has_attendance = any(keyword in question_lower for keyword in attendance_keywords)
        
        if has_attendance:
            # Tìm ngày trong câu hỏi (định dạng YYYY-MM-DD hoặc DD/MM/YYYY)
            date_patterns = [
                r'(\d{4}-\d{2}-\d{2})',  # YYYY-MM-DD
                r'(\d{1,2})/(\d{1,2})/(\d{4})',  # DD/MM/YYYY
                r'(\d{1,2})-(\d{1,2})-(\d{4})',  # DD-MM-YYYY
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, question_lower)
                if match:
                    try:
                        if pattern == r'(\d{4}-\d{2}-\d{2})':
                            query_date = datetime.strptime(match.group(0), '%Y-%m-%d').date()
                        else:
                            # DD/MM/YYYY hoặc DD-MM-YYYY
                            date_str = match.group(0).replace('-', '/')
                            query_date = datetime.strptime(date_str, '%d/%m/%Y').date()
                        
                        today = datetime.now().date()
                        
                        # Nếu ngày hỏi > ngày hôm nay, đây là ngày tương lai
                        if query_date > today:
                            return ChatResponse(
                                sql=None,
                                data=None,
                                answer=f"Ngày {query_date.strftime('%d/%m/%Y')} chưa tới, tôi không thể cung cấp dữ liệu chấm công cho ngày này.",
                                download_url=None
                            )
                    except:
                        pass  # Nếu parse lỗi, tiếp tục bình thường
        
        # BƯỚC 1: SINH SQL
        sql_chain = SQL_PROMPT | llm | StrOutputParser()
        raw_sql = sql_chain.invoke({
            "schema": HRM_SCHEMA_ENHANCED,
            "question": req.question
        })
        sql = validate_sql(raw_sql)

        # Xử lý riêng trường hợp AI trả về NO_DATA
        if "NO_DATA" in sql:
          ql = question_lower

          # 6.1. Nếu là câu hỏi về SỰ KIỆN/LỊCH TRÌNH → TỰ XÂY DỰNG SQL TRÊN BẢNG `lich_trinh`
          if any(kw in ql for kw in ["sự kiện", "su kien", "lịch trình", "lich trinh", "lịch họp", "lich hop"]):
            import re

            # Mặc định: liệt kê tất cả sự kiện, sắp xếp theo ngày bắt đầu mới nhất
            event_sql = "SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc, ngay_tao FROM lich_trinh ORDER BY ngay_bat_dau DESC"

            # Nếu câu hỏi chỉ rõ tháng/năm tạo, ví dụ: "tháng 11/2025"
            month_year_match = re.search(r"tháng\s+(\d{1,2})\s*/\s*(\d{4})", ql)
            if month_year_match:
              month = int(month_year_match.group(1))
              year = int(month_year_match.group(2))
              event_sql = (
                "SELECT tieu_de, mo_ta, ngay_bat_dau, ngay_ket_thuc, ngay_tao "
                "FROM lich_trinh "
                f"WHERE MONTH(ngay_tao) = {month} AND YEAR(ngay_tao) = {year} "
                "ORDER BY ngay_bat_dau DESC"
              )

            sql = event_sql

          else:
            # Nếu thực sự ngoài phạm vi HRM → thử trả lời từ tài liệu tĩnh (DOCX RAG)
            static_context = build_static_context(req.question, k=3)

            if static_context:
              static_chain = STATIC_ANSWER_PROMPT | llm_answer | StrOutputParser()
              static_answer = static_chain.invoke({
                "question": req.question,
                "context": static_context,
              })

              # Làm sạch Markdown in đậm nếu vẫn còn
              if isinstance(static_answer, str):
                  static_answer = static_answer.replace("**", "")

              return ChatResponse(
                sql=None,
                data=None,
                answer=static_answer,
                download_url=None,
              )

            # Nếu không có tài liệu tĩnh phù hợp → trả lời như trước
            return ChatResponse(
              sql=None,
              data=None,
              answer="Tôi chưa có dữ liệu cho nội dung bạn vừa đề cập. Hệ thống Chatbot hiện hỗ trợ tra cứu và báo cáo dữ liệu Quản lý Nhân sự (HRM), bao gồm: nhân sự, tiến độ dự án, công việc. Bạn muốn tìm hiểu thêm về chúng không?",
              download_url=None,
            )
        
        # Nếu câu hỏi liên quan đến KPI (dữ liệu không công khai)
        if "KPI_BLOCKED" in sql:
            return ChatResponse(
                sql=None,
                data=None,
                answer="Dữ liệu KPI là thông tin không công khai. Bạn có muốn hỏi về các vấn đề khác không?",
                download_url=None
            )
        
        # Nếu câu hỏi liên quan đến lịch sử nhân sự (dữ liệu không công khai)
        if "HISTORY_BLOCKED" in sql:
            return ChatResponse(
                sql=None,
                data=None,
                answer="",
                download_url=None
            )

        # BƯỚC 2: CHẠY SQL
        if not sql:
            data_result = None
            final_answer = "Xin lỗi, tôi không thể hiểu yêu cầu này."
            download_url = None
        else:
            data_result = execute_sql_api(sql)
            download_url = None
            
            # BƯỚC 2.5: VALIDATE DỮ LIỆU PHÉP NĂM (trước khi sinh câu trả lời)
            if data_result and not isinstance(data_result, str):
                leave_data_valid, leave_error_msg = validate_leave_balance_data(data_result)
                if not leave_data_valid:
                    # Nếu dữ liệu phép năm sai → Return lỗi ngay
                    return ChatResponse(
                        sql=sql,
                        data=data_result,
                        answer=leave_error_msg,
                        download_url=None
                    )
            
            # BƯỚC 3: SINH CÂU TRẢ LỜI
            if isinstance(data_result, str) and "Lỗi" in data_result:
              final_answer = f"⚠️ {data_result}"
            else:
              import json

              handled_directly = False

              # 3.1. Trường hợp đặc biệt: câu hỏi về "đi muộn" dạng liệt kê ("ai đi muộn")
              # Không áp dụng cho các câu hỏi thống kê tổng số / bao nhiêu / số lần
              is_late_question = "đi muộn" in question_lower
              is_aggregate_late = any(
                kw in question_lower
                for kw in [
                  "tổng",
                  "tổng số",
                  "bao nhiêu",
                  "bao nhieu",
                  "số lần",
                  "so lan",
                  "lần",
                  "lan",
                ]
              )

              if is_late_question and not is_aggregate_late:
                rows = []
                if isinstance(data_result, dict) and 'data' in data_result:
                  rows = data_result.get('data') or []
                elif isinstance(data_result, list):
                  rows = data_result

                if isinstance(rows, list) and rows:
                  # Chỉ in danh sách từng nhân viên, không cần dòng tiêu đề tổng quát
                  lines = []
                  for record in rows:
                    if isinstance(record, dict):
                      name = record.get('ho_ten') or record.get('ten_nhan_vien') or record.get('nhan_vien') or ""
                      emp_id = record.get('nhan_vien_id') or record.get('id_nhan_vien')
                      date = record.get('ngay_cham_cong') or record.get('ngay') or ""
                      check_in = record.get('check_in') or record.get('gio_vao') or ""
                      check_out = record.get('check_out') or record.get('gio_ra') or ""

                      parts = []
                      if name:
                        if emp_id is not None:
                          parts.append(f"{name} (ID: {emp_id})")
                        else:
                          parts.append(name)
                      elif emp_id is not None:
                        parts.append(f"Nhân viên ID {emp_id}")

                      if date:
                        parts.append(f"ngày {date}")
                      if check_in:
                        parts.append(f"check in {check_in}")
                      if check_out:
                        parts.append(f"check out {check_out}")

                      if parts:
                        line = "- " + ", ".join(parts)
                      else:
                        line = "- " + json.dumps(record, ensure_ascii=False)
                    else:
                      line = "- " + str(record)
                    lines.append(line)

                  final_answer = "\n".join(lines)
                  handled_directly = True

              # 3.2. Trường hợp dữ liệu thống kê COUNT/SUM/AVG đơn giản → trả câu tự nhiên, không in JSON
              if not handled_directly:
                rows = None
                if isinstance(data_result, dict) and 'data' in data_result:
                  rows = data_result.get('data')
                elif isinstance(data_result, list):
                  rows = data_result

                if isinstance(rows, list) and len(rows) == 1 and isinstance(rows[0], dict):
                  record = rows[0]
                  # Nếu chỉ có 1 field numeric → coi là kết quả thống kê đơn giản
                  numeric_items = [
                    (k, v) for k, v in record.items()
                    if isinstance(v, (int, float))
                  ]
                  if len(numeric_items) == 1:
                    key, value = numeric_items[0]
                    # Trả câu tự nhiên hơn cho các câu hỏi thống kê số lượng
                    ql = question_lower
                    object_label = "mục"
                    # Ưu tiên loại đối tượng được hỏi đến trực tiếp
                    # Nếu câu hỏi nói về "bước" (các bước trong quy trình) → dùng "bước"
                    if "bước" in ql or "buoc" in ql:
                      object_label = "bước"
                    # Nếu câu hỏi nói về "công việc" → dùng "công việc"
                    elif "công việc" in ql or "cong viec" in ql:
                      object_label = "công việc"
                    # Nếu câu hỏi nói về "nhân viên" → dùng "nhân viên"
                    elif "nhân viên" in ql or "nhan vien" in ql:
                      object_label = "nhân viên"
                    # Nếu câu hỏi nói về "dự án" → dùng "dự án"
                    elif "dự án" in ql or "du an" in ql:
                      object_label = "dự án"
                    elif "tài liệu" in ql or "tai lieu" in ql:
                      object_label = "tài liệu"
                    elif "sự kiện" in ql or "su kien" in ql:
                      object_label = "sự kiện"
                    elif "phòng ban" in ql or "phong ban" in ql:
                      object_label = "phòng ban"
                    elif "quyền" in ql or "quyen" in ql:
                      object_label = "quyền"

                    # Nếu câu hỏi nói về "số lần" thì label là "lần"
                    if "lần" in ql or "lan" in ql:
                      object_label = "lần"

                    # Nếu record có thêm tên công việc/dự án, ưu tiên trả lời kèm tên
                    name_field = record.get("ten_cong_viec") or record.get("ten_du_an") or record.get("ten")
                    question_text = req.question.strip()
                    q_lower_full = question_text.lower()

                    # 1) Câu hỏi dạng "Có bao nhiêu ...?" → "Có X ..."
                    if q_lower_full.startswith("có bao nhiêu") or q_lower_full.startswith("co bao nhieu"):
                      if q_lower_full.startswith("có bao nhiêu"):
                        prefix_len = len("có bao nhiêu")
                      else:
                        prefix_len = len("co bao nhieu")

                      suffix = question_text[prefix_len:].lstrip(" ,.:-")
                      suffix = suffix.rstrip(" ?!.")
                      if suffix:
                        final_answer = f"Có {value} {suffix}."
                      else:
                        final_answer = f"Có {value} {object_label}."

                    # 2) Câu hỏi dạng "Tổng số ..." → "Tổng số ... là X <label>."
                    elif q_lower_full.startswith("tổng số") or q_lower_full.startswith("tong so"):
                      base_q = question_text.rstrip(" ?!.")
                      final_answer = f"{base_q} là {value} {object_label}."

                    # 3) Câu hỏi dạng "X nào ... nhiều nhất" với tên + số lượng → nêu rõ tên
                    elif ("nào" in ql or "nao" in ql) and name_field and ("nhiều nhất" in ql or "nhieu nhat" in ql):
                      # Ví dụ: "Công việc nào có nhân viên thực hiện nhiều nhất?"
                      final_answer = f"Công việc {name_field} có {value} {object_label}."

                    # 4) Mặc định: giữ câu ngắn gọn
                    else:
                      final_answer = f"Có {value} {object_label}."
                    handled_directly = True

              # 3.3. Mặc định: gửi cho LLM format theo ANSWER_PROMPT
              if not handled_directly:
                # Convert data thành JSON string rõ ràng để LLM dễ parse
                if data_result is None:
                  data_str = "null"
                elif isinstance(data_result, str):
                  data_str = data_result
                elif isinstance(data_result, dict) and 'data' in data_result:
                  data_str = json.dumps(data_result['data'], ensure_ascii=False)
                else:
                  data_str = json.dumps(data_result, ensure_ascii=False)

                ans_chain = ANSWER_PROMPT | llm_answer | StrOutputParser()
                final_answer = ans_chain.invoke({
                  "question": req.question,
                  "data": data_str
                })

                # HẬU KIỂM ĐẶC BIỆT CHO CÂU HỎI "TRỄ HẠN":
                # Nếu SQL đã trả về dữ liệu nhưng LLM vẫn trả lời "Không có dự án nào trễ hạn"
                # thì override bằng cách liệt kê trực tiếp từ data_result.
                if ("trễ hạn" in question_lower or "tre han" in question_lower) and data_result and not isinstance(data_result, str):
                  rows = []
                  if isinstance(data_result, dict) and 'data' in data_result:
                    rows = data_result.get('data') or []
                  elif isinstance(data_result, list):
                    rows = data_result

                  if isinstance(rows, list) and rows:
                    text_lower = final_answer.lower() if isinstance(final_answer, str) else ""
                    if "không có dự án nào trễ hạn" in text_lower:
                      lines = []
                      for record in rows:
                        if isinstance(record, dict):
                          name = record.get('ten_du_an') or record.get('ten_cong_viec') or record.get('ten') or record.get('tieu_de') or ""
                          end_date = record.get('ngay_ket_thuc') or record.get('han_hoan_thanh') or ""
                          status = record.get('trang_thai_duan') or record.get('trang_thai') or ""

                          parts = []
                          if name:
                            parts.append(name)
                          if end_date:
                            parts.append(f"ngày kết thúc {end_date}")
                          if status:
                            parts.append(f"trạng thái {status}")

                          if parts:
                            line = "- " + ", ".join(parts)
                          else:
                            line = "- " + json.dumps(record, ensure_ascii=False)
                        else:
                          line = "- " + str(record)
                        lines.append(line)

                      final_answer = "\n".join(lines)
            
            # BƯỚC 4: KIỂM TRA YÊU CẦU XUẤT FILE (sau khi có câu trả lời)
            q_lower = req.question.lower()
            
            if data_result and not isinstance(data_result, str):
                # Nếu có dữ liệu và người dùng yêu cầu xuất file
                if "word" in q_lower or "docx" in q_lower or "văn bản" in q_lower or "xuất" in q_lower or "file" in q_lower:
                    try:
                        file_path = create_word_report(
                            data=data_result, 
                            title="BÁO CÁO TRUY VẤN HRM", 
                            filename_prefix="baocao",
                            question=req.question,
                            summary=final_answer
                        )
                        if file_path:
                            # Lấy tên file từ path
                            filename = os.path.basename(file_path)
                            download_url = f"/download/{filename}"
                    except Exception as e:
                        print(f"Error creating word report: {e}")

        return ChatResponse(
            sql=sql,
            data=data_result,
            answer=final_answer,
            download_url=download_url
        )

    except Exception as e:
        print(f"Server Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))